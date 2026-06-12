"""
compare_versions.py - 对比两次PD主攻词单，生成词库更新报告

Workflow:
  1. --auto 或手动 --prev/--curr → compare → stats JSON
  2. Agent 写 insights_tracker.md
  3. --insights-file 重建 Word
"""
import argparse
import json
import os
import sys
import warnings
from pathlib import Path

warnings.filterwarnings("ignore")

SCOUT_SCRIPTS = Path(__file__).resolve().parents[2] / "sif-keyword-scout" / "scripts"
sys.path.insert(0, str(SCOUT_SCRIPTS))

from check_history import check_history, pd_excel_path, find_compare_baseline
from report_utils import (
    get_col, load_insights, save_stats, cleanup_intermediate_files,
    setup_doc_styles, add_doc_title, add_para, set_cell_text,
    apply_font_to_table, finalize_doc_fonts, render_insights_section,
    insights_path, llm_prompt_template, is_missing, fmt_share_pct, fmt_number,
)

try:
    import pandas as pd
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Border, Side
except ImportError as e:
    print(f"缺少依赖：{e}", file=sys.stderr)
    sys.exit(1)

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

THIN = Side(style="thin")
THIN_BORDER = Border(left=THIN, right=THIN, top=THIN, bottom=THIN)


def load_pd_sheet(path: str) -> pd.DataFrame:
    xl = pd.ExcelFile(path)
    dfs = []
    for sheet in xl.sheet_names:
        if "SSS" in sheet or "SS" in sheet or "S级" in sheet or "PD" in sheet:
            try:
                dfs.append(pd.read_excel(path, sheet_name=sheet, header=0))
            except Exception:
                pass
    if not dfs:
        return pd.read_excel(path, sheet_name=0, header=0)
    return pd.concat(dfs, ignore_index=True)


def normalize_kw(kw) -> str:
    return str(kw).lower().strip()


def compare(prev_df: pd.DataFrame, curr_df: pd.DataFrame) -> dict:
    col_kw_prev = get_col(prev_df, ["关键词", "Keyword"])
    col_kw_curr = get_col(curr_df, ["关键词", "Keyword"])
    col_lv_prev = get_col(prev_df, ["词级别", "Level"])
    col_lv_curr = get_col(curr_df, ["词级别", "Level"])
    col_sv_curr = get_col(curr_df, ["周搜索量", "搜索量"])
    col_sv_prev = get_col(prev_df, ["周搜索量", "搜索量"])
    col_sh_curr = get_col(curr_df, ["竞品SP份额%"])
    col_sh_prev = get_col(prev_df, ["竞品SP份额%"])
    col_src_curr = get_col(curr_df, ["出现来源", "来源"])

    if not col_kw_prev or not col_kw_curr:
        raise ValueError("找不到关键词列，请检查PD主攻词单格式")

    prev_df = prev_df.copy()
    curr_df = curr_df.copy()
    prev_df["__kw"] = prev_df[col_kw_prev].apply(normalize_kw)
    curr_df["__kw"] = curr_df[col_kw_curr].apply(normalize_kw)

    prev_dict = prev_df.set_index("__kw").to_dict("index")
    curr_dict = curr_df.set_index("__kw").to_dict("index")

    prev_set = set(prev_dict.keys()) - {"", "nan"}
    curr_set = set(curr_dict.keys()) - {"", "nan"}

    new_words = curr_set - prev_set
    removed_words = prev_set - curr_set
    common_words = prev_set & curr_set

    def get_grade(d, kw, col_lv):
        return str(d.get(kw, {}).get(col_lv, "")) if col_lv else ""

    def get_sv(d, kw, col_sv):
        if not col_sv:
            return None
        raw = d.get(kw, {}).get(col_sv)
        if is_missing(raw):
            return None
        try:
            return float(raw)
        except Exception:
            return None

    def get_sh(d, kw, col_sh):
        if not col_sh:
            return None
        raw = d.get(kw, {}).get(col_sh)
        if is_missing(raw):
            return None
        try:
            return float(raw)
        except Exception:
            return None

    new_by_grade = {"SSS级": [], "SS级": [], "S级": [], "其他": []}
    for kw in new_words:
        grade = get_grade(curr_dict, kw, col_lv_curr)
        key = grade if grade in new_by_grade else "其他"
        new_by_grade[key].append({
            "关键词": curr_dict[kw].get(col_kw_curr, kw),
            "级别": grade,
            "周搜索量": get_sv(curr_dict, kw, col_sv_curr),
            "来源": str(curr_dict[kw].get(col_src_curr, "")) if col_src_curr else "",
        })

    removed_by_grade = {"SSS级": [], "SS级": [], "S级": [], "其他": []}
    for kw in removed_words:
        grade = get_grade(prev_dict, kw, col_lv_prev)
        key = grade if grade in removed_by_grade else "其他"
        removed_by_grade[key].append({
            "关键词": prev_dict[kw].get(col_kw_prev, kw),
            "原级别": grade,
        })

    sv_changed = []
    for kw in common_words:
        sv_prev = get_sv(prev_dict, kw, col_sv_prev)
        sv_curr = get_sv(curr_dict, kw, col_sv_curr)
        if sv_prev is None or sv_curr is None or sv_prev <= 0:
            continue
        change_pct = (sv_curr - sv_prev) / sv_prev * 100
        if abs(change_pct) > 20:
            sv_changed.append({
                "关键词": curr_dict[kw].get(col_kw_curr, kw),
                "上次搜索量": int(sv_prev),
                "本次搜索量": int(sv_curr),
                "变化%": round(change_pct, 1),
                "方向": "↑涨" if change_pct > 0 else "↓跌",
                "级别": get_grade(curr_dict, kw, col_lv_curr),
            })
    sv_changed.sort(key=lambda x: abs(x["变化%"]), reverse=True)

    sh_changed = []
    for kw in common_words:
        sh_prev = get_sh(prev_dict, kw, col_sh_prev)
        sh_curr = get_sh(curr_dict, kw, col_sh_curr)
        if sh_prev is None and sh_curr is None:
            continue
        prev_v = sh_prev if sh_prev is not None else 0.0
        curr_v = sh_curr if sh_curr is not None else 0.0
        if abs(curr_v - prev_v) > 5:
            sh_changed.append({
                "关键词": curr_dict[kw].get(col_kw_curr, kw),
                "上次份额%": None if sh_prev is None else round(sh_prev, 2),
                "本次份额%": None if sh_curr is None else round(sh_curr, 2),
                "变化": round(curr_v - prev_v, 2),
            })

    prev_sss = {kw for kw in prev_set if get_grade(prev_dict, kw, col_lv_prev) == "SSS级"}
    curr_sss = {kw for kw in curr_set if get_grade(curr_dict, kw, col_lv_curr) == "SSS级"}
    prev_ss = {kw for kw in prev_set if get_grade(prev_dict, kw, col_lv_prev) == "SS级"}
    curr_ss = {kw for kw in curr_set if get_grade(curr_dict, kw, col_lv_curr) == "SS级"}

    ss_stable = []
    for kw in prev_ss & curr_ss:
        ss_stable.append(curr_dict[kw].get(col_kw_curr, kw))

    return {
        "summary": {
            "prev_total": len(prev_set),
            "curr_total": len(curr_set),
            "new_total": len(new_words),
            "removed_total": len(removed_words),
            "sv_changed_total": len(sv_changed),
            "sh_changed_total": len(sh_changed),
            "sss_new": len(curr_sss - prev_sss),
            "sss_removed": len(prev_sss - curr_sss),
            "sss_stable": len(prev_sss & curr_sss),
            "ss_new": len(curr_ss - prev_ss),
            "ss_removed": len(prev_ss - curr_ss),
            "ss_stable": len(prev_ss & curr_ss),
        },
        "new_by_grade": new_by_grade,
        "removed_by_grade": removed_by_grade,
        "sv_changed": sv_changed[:20],
        "sh_changed": sh_changed[:15],
        "sss_new_list": [curr_dict[kw].get(col_kw_curr, kw) for kw in curr_sss - prev_sss],
        "sss_removed_list": [prev_dict[kw].get(col_kw_prev, kw) for kw in prev_sss - curr_sss],
        "sss_stable_list": [curr_dict[kw].get(col_kw_curr, kw) for kw in prev_sss & curr_sss],
        "ss_new_list": [curr_dict[kw].get(col_kw_curr, kw) for kw in curr_ss - prev_ss],
        "ss_removed_list": [prev_dict[kw].get(col_kw_prev, kw) for kw in prev_ss - curr_ss],
        "ss_stable_list": ss_stable[:20],
    }


def _add_word_table(doc, headers, rows):
    if not rows:
        return
    tbl = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    tbl.style = "Table Grid"
    for ci, h in enumerate(headers):
        set_cell_text(tbl.rows[0].cells[ci], h)
    for ri, row in enumerate(rows, 1):
        for ci, val in enumerate(row):
            if isinstance(val, float) and pd.isna(val):
                val = ""
            elif val is None:
                val = ""
            set_cell_text(tbl.rows[ri].cells[ci], fmt_cell(val))
    apply_font_to_table(tbl)
    doc.add_paragraph()


def _grade_word_list(doc, words: list, label: str):
    if not words:
        add_para(doc, f"{label}：无")
        return
    add_para(doc, f"{label}（{len(words)}个）：", bold=True)
    for w in words:
        if isinstance(w, dict):
            sv = w.get("周搜索量")
            sv_text = fmt_number(sv) if sv is not None else ""
            src = w.get("来源") or w.get("原级别", "")
            extra = f"，搜索量 {sv_text}" if sv_text else ""
            if src:
                extra += f"，{src}"
            add_para(doc, f"  • {w.get('关键词', w)}{extra}")
        else:
            add_para(doc, f"  • {w}")


def build_word_report(compare_result, asin, prev_date, curr_date, output_dir,
                      insights="", compare_note="", days_apart=None):
    if not HAS_DOCX:
        print("WARN: no python-docx", file=sys.stderr)
        return None

    doc = Document()
    setup_doc_styles(doc)
    days_str = f"（相隔 {days_apart} 天）" if days_apart is not None else ""
    add_doc_title(doc, f"{asin} 词库更新报告",
                  f"对比：{prev_date} → {curr_date}{days_str}  |  {compare_note}")

    s = compare_result["summary"]
    new_bg = compare_result["new_by_grade"]
    rem_bg = compare_result["removed_by_grade"]

    doc.add_heading("一、变化摘要（数据）", level=1)
    add_para(doc,
        f"上次词库共 {s['prev_total']} 个词，本次共 {s['curr_total']} 个词。"
        f"新增 {s['new_total']} 个（SSS:{len(new_bg['SSS级'])} SS:{len(new_bg['SS级'])} S:{len(new_bg['S级'])}），"
        f"消失 {s['removed_total']} 个（SSS:{len(rem_bg['SSS级'])} SS:{len(rem_bg['SS级'])} S:{len(rem_bg['S级'])}）。"
        f"搜索量变化>20%：{s['sv_changed_total']} 个；竞品份额变化>5%：{s['sh_changed_total']} 个。"
        f"SS 级稳定 {s['ss_stable']} 个，SS 新增 {s['ss_new']} 个，SS 消失 {s['ss_removed']} 个。"
    )

    doc.add_heading("二、策略解读与投放调整（AI 分析）", level=1)
    render_insights_section(doc, insights, heading="")

    doc.add_heading("三、新增词清单", level=1)
    has_new = False
    for grade in ["SSS级", "SS级", "S级"]:
        words = new_bg.get(grade, [])
        if words:
            has_new = True
            sorted_w = sorted(words, key=lambda x: x.get("周搜索量", 0) or 0, reverse=True)
            _grade_word_list(doc, sorted_w, grade)
    if not has_new:
        add_para(doc, "本次无新增词。")

    doc.add_heading("四、消失词清单", level=1)
    has_rem = False
    for grade in ["SSS级", "SS级", "S级"]:
        words = rem_bg.get(grade, [])
        if words:
            has_rem = True
            _grade_word_list(doc, words, f"原{grade}消失")
    if not has_rem:
        add_para(doc, "本次无消失词。")

    doc.add_heading("五、SS/SSS 级稳定词", level=1)
    if compare_result["ss_stable_list"] or compare_result["sss_stable_list"]:
        if compare_result["sss_stable_list"]:
            _grade_word_list(doc, compare_result["sss_stable_list"], "稳定 SSS 级")
        if compare_result["ss_stable_list"]:
            _grade_word_list(doc, compare_result["ss_stable_list"], "稳定 SS 级")
    else:
        add_para(doc, "无跨版本稳定的 SS/SSS 级词（或两级均为 0）。")

    doc.add_heading("六、搜索量变化 > 20%", level=1)
    sv_ch = compare_result["sv_changed"]
    if sv_ch:
        rows = [[r["关键词"], r["级别"], r["上次搜索量"], r["本次搜索量"],
                 f"{r['变化%']}%", r["方向"]] for r in sv_ch[:15]]
        _add_word_table(doc, ["关键词", "级别", "上次", "本次", "变化", "方向"], rows)
    else:
        add_para(doc, "无搜索量变化超过 20% 的词。")

    doc.add_heading("七、竞品 SP 份额变化 > 5%", level=1)
    sh_ch = compare_result["sh_changed"]
    if sh_ch:
        rows = [[r["关键词"], r["上次份额%"], r["本次份额%"], r["变化"]] for r in sh_ch[:10]]
        _add_word_table(doc, ["关键词", "上次份额%", "本次份额%", "变化"], rows)
    else:
        add_para(doc, "无竞品 SP 份额变化超过 5% 的词。")

    finalize_doc_fonts(doc)
    out = os.path.join(output_dir, f"{asin}_词库更新报告_{curr_date}.docx")
    try:
        doc.save(out)
    except PermissionError:
        alt = os.path.join(output_dir, f"{asin}_词库更新报告_{curr_date}_new.docx")
        doc.save(alt)
        out = alt
    print(f"OK tracker report: {out}")
    return out


def resolve_auto_paths(output_root: str, asin: str, curr_date: str, compare_window: int):
    hist = check_history(asin, output_root, curr_date=curr_date, compare_window=compare_window)
    if not hist.get("compare_prev_result_dir"):
        print("ERROR: no compare baseline in history", file=sys.stderr)
        sys.exit(1)
    prev_date = hist["compare_prev_date"]
    prev_xlsx = pd_excel_path(hist["compare_prev_result_dir"], asin, prev_date)
    curr_xlsx = pd_excel_path(
        os.path.join(output_root, asin, curr_date, "处理结果"), asin, curr_date
    )
    for p in (prev_xlsx, curr_xlsx):
        if not os.path.exists(p):
            print(f"ERROR: missing {p}", file=sys.stderr)
            sys.exit(1)
    return prev_xlsx, curr_xlsx, prev_date, hist


def main():
    parser = argparse.ArgumentParser(description="对比两次PD主攻词单")
    parser.add_argument("--prev", default="")
    parser.add_argument("--curr", default="")
    parser.add_argument("--output-dir", required=True, help="本次处理结果目录")
    parser.add_argument("--output-root", default="", help="ASIN 根目录（--auto 时用）")
    parser.add_argument("--asin", required=True)
    parser.add_argument("--prev-date", default="")
    parser.add_argument("--curr-date", required=True)
    parser.add_argument("--auto", action="store_true", help="从历史表按 1~7 天窗口自动选对比基准")
    parser.add_argument("--compare-window", type=int, default=7)
    parser.add_argument("--insights", default="")
    parser.add_argument("--insights-file", default="")
    parser.add_argument("--skip-word", action="store_true")
    parser.add_argument("--dashboard", default="")
    args = parser.parse_args()
    args.insights = load_insights(args.insights, args.insights_file)

    output_root = args.output_root or str(Path(args.output_dir).parents[2])
    compare_note = ""
    days_apart = None

    if args.auto:
        prev_path, curr_path, prev_date, hist = resolve_auto_paths(
            output_root, args.asin, args.curr_date, args.compare_window
        )
        args.prev = prev_path
        args.curr = curr_path
        args.prev_date = prev_date
        compare_note = hist.get("compare_note", "")
        days_apart = hist.get("compare_days_apart")
    else:
        if not args.prev or not args.curr or not args.prev_date:
            print("ERROR: need --prev --curr --prev-date or use --auto", file=sys.stderr)
            sys.exit(1)
        prev_path, curr_path = args.prev, args.curr

    for p in (args.prev, args.curr):
        if not os.path.exists(p):
            print(f"ERROR: not found {p}", file=sys.stderr)
            sys.exit(1)

    os.makedirs(args.output_dir, exist_ok=True)

    print(f"prev: {args.prev}")
    print(f"curr: {args.curr}")
    prev_df = load_pd_sheet(args.prev)
    curr_df = load_pd_sheet(args.curr)
    compare_result = compare(prev_df, curr_df)

    stats = {
        "report": "tracker_compare",
        "asin": args.asin,
        "prev_date": args.prev_date,
        "curr_date": args.curr_date,
        "compare_note": compare_note,
        "compare_days_apart": days_apart,
        **compare_result,
    }
    stats_path = os.path.join(args.output_dir, f"{args.asin}_tracker_stats_{args.curr_date}.json")
    save_stats(stats, stats_path)
    print(f"stats: {stats_path}")
    print(json.dumps(compare_result["summary"], ensure_ascii=False, indent=2))

    if args.skip_word:
        ip = os.path.join(args.output_dir, "insights_tracker.md")
        print("=" * 60)
        print(f"[Agent] write {ip} then --insights-file rebuild Word")
        print(llm_prompt_template("tracker", stats))
        print("=" * 60)
        return

    build_word_report(compare_result, args.asin, args.prev_date, args.curr_date,
                      args.output_dir, insights=args.insights,
                      compare_note=compare_note, days_apart=days_apart)
    cleanup_intermediate_files(args.output_dir, args.asin, args.curr_date, "tracker")


if __name__ == "__main__":
    main()
