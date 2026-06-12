"""
report_utils.py - 报告生成共享工具
- stats JSON 供 Agent 写分析段
- 图表 PNG → Word 内嵌（不保留 charts/）
- 最终产物：仅 Excel + Word (.docx)
"""
from __future__ import annotations

import stdio_utf8  # noqa: F401 — Windows GBK 控制台 UTF-8 兼容

import json
import math
import os
import re
import shutil
from datetime import datetime
from pathlib import Path

FONT_NAME = "Microsoft YaHei"


# ─── 路径工具 ───────────────────────────────────────────────

def charts_dir(output_dir: str) -> str:
    d = os.path.join(output_dir, "charts")
    os.makedirs(d, exist_ok=True)
    return d


def chart_path(output_dir: str, name: str) -> str:
    return os.path.join(charts_dir(output_dir), name)


def get_col(df, candidates: list):
    """按候选列名列表找第一个匹配列（精确优先，模糊兜底）。"""
    cols = df.columns.tolist()
    for c in candidates:
        if c in cols:
            return c
    for c in candidates:
        for col in cols:
            if c in str(col):
                return col
    return None


def is_missing(val) -> bool:
    """空值、NaN、无效占位 → 视为缺失（展示用空白）。"""
    if val is None:
        return True
    if isinstance(val, float) and (math.isnan(val) or math.isinf(val)):
        return True
    try:
        if pd_is_na(val):
            return True
    except Exception:
        pass
    s = str(val).strip().lower()
    return s in ("", "nan", "none", "nat", "-")


def pd_is_na(val) -> bool:
    try:
        import pandas as pd
        return pd.isna(val)
    except Exception:
        return False


def fmt_cell(val) -> str:
    """Excel/Word 单元格：缺失 → 空字符串。"""
    if is_missing(val):
        return ""
    if isinstance(val, float) and val == int(val):
        return str(int(val))
    return str(val).strip()


def fmt_share_pct(val) -> str:
    """份额列：缺失 → 空白；有值 → x.xx%"""
    if is_missing(val):
        return ""
    return f"{pct_to_float(val):.2f}%"


def fmt_number(val, decimals: int = 0) -> str:
    if is_missing(val):
        return ""
    try:
        v = float(val)
        if decimals <= 0:
            return str(int(v))
        return f"{v:.{decimals}f}"
    except (ValueError, TypeError):
        return fmt_cell(val)


def _sanitize_for_json(obj):
    if isinstance(obj, dict):
        return {k: _sanitize_for_json(v) for k, v in obj.items()}
    if isinstance(obj, list):
        return [_sanitize_for_json(v) for v in obj]
    if isinstance(obj, float) and (math.isnan(obj) or math.isinf(obj)):
        return None
    return obj


def pct_to_float(val) -> float:
    """将百分比统一转为 0~100 数值（Sif 导出可能是 '18.21%' 或 0.1821 小数）。"""
    try:
        if val is None:
            return 0.0
        if isinstance(val, float) and val != val:
            return 0.0
        raw = str(val).strip()
        if not raw or raw.lower() in ("nan", "none", "-", ""):
            return 0.0
        has_pct = raw.endswith("%")
        v = float(raw.replace("%", "").strip())
        if has_pct or v > 1.0:
            return v
        return v * 100.0
    except (ValueError, TypeError):
        return 0.0


def stats_path(output_dir: str, asin: str, report_key: str, run_id: str) -> str:
    return os.path.join(output_dir, f"{asin}_{report_key}_stats_{run_id}.json")


def insights_path(output_dir: str, report_key: str) -> str:
    return os.path.join(output_dir, f"insights_{report_key}.md")


# ─── stats / insights ───────────────────────────────────────

def load_insights(insights: str = "", insights_file: str = "") -> str:
    if insights_file and os.path.exists(insights_file):
        with open(insights_file, encoding="utf-8") as f:
            return f.read().strip()
    return (insights or "").strip()


def save_stats(stats: dict, path: str) -> str:
    stats["generated_at"] = datetime.now().isoformat(timespec="seconds")
    clean = _sanitize_for_json(stats)
    with open(path, "w", encoding="utf-8") as f:
        json.dump(clean, f, ensure_ascii=False, indent=2)
    return path


def cleanup_intermediate_files(output_dir: str, asin: str = "", run_id: str = "", report_key: str = ""):
    """Word 生成后清理：只保留 xlsx + docx。仅删除本报告对应的中间文件。"""
    removed = []
    p = Path(output_dir)
    if not p.is_dir():
        return removed

    for f in p.iterdir():
        if not f.is_file():
            continue
        suffix = f.suffix.lower()
        name = f.name
        if report_key:
            if suffix == ".json" and name in (
                f"{asin}_{report_key}_stats_{run_id}.json",
                f"{asin}_tracker_stats_{run_id}.json",
            ):
                f.unlink(missing_ok=True)
                removed.append(name)
            elif suffix == ".md" and name == f"insights_{report_key}.md":
                f.unlink(missing_ok=True)
                removed.append(name)
        elif suffix == ".json" and asin and name.startswith(asin) and name.endswith(".json"):
            f.unlink(missing_ok=True)
            removed.append(name)

    charts = p / "charts"
    if charts.is_dir():
        shutil.rmtree(charts, ignore_errors=True)
        removed.append("charts/")

    if removed:
        print(f"🧹 已清理中间文件：{', '.join(removed)}")
    return removed


# ─── Word 样式与段落 ────────────────────────────────────────

def _set_run_font(run, size: int = 11, bold: bool = False, italic: bool = False, color=None):
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), FONT_NAME)


def _shade_cell(cell, fill_hex: str = "1F4E79"):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def setup_doc_styles(doc):
    from docx.shared import Pt, Cm, RGBColor
    from docx.oxml.ns import qn
    sec = doc.sections[0]
    sec.top_margin = Cm(2.2)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)

    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)

    for lvl, size, color in [(1, 15, "1F4E79"), (2, 12, "2E75B6"), (3, 11, "404040")]:
        h = doc.styles[f"Heading {lvl}"]
        h.font.name = FONT_NAME
        h.font.size = Pt(size)
        h.font.bold = True
        h.font.color.rgb = RGBColor(int(color[:2], 16), int(color[2:4], 16), int(color[4:], 16))
        h._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)

    for style_name in ("List Bullet", "List Number"):
        if style_name in doc.styles:
            s = doc.styles[style_name]
            s.font.name = FONT_NAME
            s.font.size = Pt(11)
            s._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)


def add_para(doc, text: str, size: int = 11, bold: bool = False, space_after: int = 6):
    from docx.shared import Pt
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    _set_run_font(run, size=size, bold=bold)
    return p


def add_doc_title(doc, title: str, subtitle: str = ""):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    _set_run_font(run, size=18, bold=True, color=(0x1F, 0x4E, 0x79))
    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(subtitle)
        _set_run_font(run2, size=10, color=(0x66, 0x66, 0x66))
    doc.add_paragraph()


def set_cell_text(cell, text: str, size: int = 10, bold: bool = False):
    """写入单元格并强制微软雅黑（避免 .text= 丢失字体）"""
    if text is None or str(text).strip().lower() in ("nan", "none", "nat"):
        text = ""
    cell.text = ""
    p = cell.paragraphs[0]
    _add_rich_runs(p, str(text), size=size)
    if bold:
        for run in p.runs:
            run.font.bold = True


def apply_font_to_table(table, header_rows: int = 1):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import RGBColor
    for ri, row in enumerate(table.rows):
        for cell in row.cells:
            for para in cell.paragraphs:
                if ri < header_rows:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    _set_run_font(run, size=10, bold=(ri < header_rows))
                    if ri < header_rows:
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            if ri < header_rows:
                _shade_cell(cell)


def finalize_doc_fonts(doc):
    """遍历全文所有 w:r，强制微软雅黑（含标题、表格、列表）"""
    from docx.text.paragraph import Paragraph

    heading_sizes = {"Heading 1": 15, "Heading 2": 12, "Heading 3": 11, "Title": 18}

    def _fix_paragraph(p: Paragraph):
        style_name = p.style.name if p.style else ""
        default_size = heading_sizes.get(style_name, 11)
        for run in p.runs:
            if style_name in heading_sizes:
                size = heading_sizes[style_name]
            else:
                size = int(run.font.size.pt) if run.font.size else default_size
            _set_run_font(
                run,
                size=size,
                bold=bool(run.font.bold) or style_name.startswith("Heading"),
                italic=bool(run.font.italic),
            )

    for para in doc.paragraphs:
        _fix_paragraph(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _fix_paragraph(para)

    # XML 层兜底：覆盖 Heading 等样式未落到 run 的情况
    try:
        from docx.oxml import OxmlElement
        body = doc.element.body
        for r_elem in body.iter(qn("w:r")):
            rPr = r_elem.get_or_add_rPr()
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = OxmlElement("w:rFonts")
                rPr.insert(0, rFonts)
            rFonts.set(qn("w:ascii"), FONT_NAME)
            rFonts.set(qn("w:hAnsi"), FONT_NAME)
            rFonts.set(qn("w:eastAsia"), FONT_NAME)
            rFonts.set(qn("w:cs"), FONT_NAME)
    except Exception:
        pass


def add_picture_centered(doc, path: str, width_inches: float = 5.5):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches
    if not os.path.isfile(path):
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Inches(width_inches))


def add_charts_section(doc, chart_items: list[tuple[str, str]], section_title: str = "可视化图表", level: int = 1):
    """chart_items: [(图注说明, png路径), ...]"""
    if not chart_items:
        return
    doc.add_heading(section_title, level=level)
    add_para(doc, "说明：Word 内嵌图与 Excel 看板图表数据一致；柱形图看数量/占比，散点图看两维分布关系。", size=10)
    for idx, (caption, path) in enumerate(chart_items, 1):
        if not path or not os.path.isfile(path):
            continue
        add_para(doc, caption, size=11, bold=True, space_after=4)
        add_picture_centered(doc, path)
        doc.add_paragraph()


# ─── Markdown → Word 渲染（AI 分析段）──────────────────────

_IMG_RE = re.compile(r"!\[.*?\]\(.*?\)")
_HRULE_RE = re.compile(r"^---+\s*$")


def _parse_table_row(line: str) -> list[str]:
    line = line.strip().strip("|")
    return [c.strip() for c in line.split("|")]


def _is_table_sep(line: str) -> bool:
    return bool(re.match(r"^\|?[\s\-:|]+\|?$", line.strip()))


def _add_rich_runs(paragraph, text: str, size: int = 11):
    """解析 **bold** / ***bold-italic***，写入 Word runs"""
    pattern = re.compile(r"(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|`(.+?)`)")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            run = paragraph.add_run(text[pos:m.start()])
            _set_run_font(run, size=size)
        if m.group(2):
            run = paragraph.add_run(m.group(2))
            _set_run_font(run, size=size, bold=True, italic=True)
        elif m.group(3):
            run = paragraph.add_run(m.group(3))
            _set_run_font(run, size=size, bold=True)
        elif m.group(4):
            run = paragraph.add_run(m.group(4))
            _set_run_font(run, size=size)
        pos = m.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        _set_run_font(run, size=size)


def _add_word_table(doc, rows: list[list[str]]):
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=ncols)
    tbl.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci in range(ncols):
            cell_text = row[ci] if ci < len(row) else ""
            cell = tbl.rows[ri].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            _add_rich_runs(p, cell_text, size=10)
    apply_font_to_table(tbl, header_rows=1)
    doc.add_paragraph()


def render_insights_section(doc, insights: str, heading: str = "策略洞察"):
    """
    将 Agent 写的 Markdown 分析段渲染进 Word。
    支持：##/### 标题、- 列表、1. 编号列表、|表格|、**加粗**
    自动跳过：![图片](...) 、--- 分隔线
    """
    if heading:
        doc.add_heading(heading, level=1)

    if not insights or not insights.strip():
        add_para(doc, "（暂无 AI 分析，请 Agent 根据 stats JSON 撰写 insights 后重建 Word）", size=10)
        return

    lines = insights.splitlines()
    i = 0
    table_buf: list[list[str]] = []

    def flush_table():
        nonlocal table_buf
        if table_buf:
            _add_word_table(doc, table_buf)
            table_buf = []

    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()

        # 跳过图片语法（图表由 Word 构建函数单独插入）
        if _IMG_RE.search(line) or line.strip().startswith("!["):
            i += 1
            continue
        if _HRULE_RE.match(line.strip()):
            i += 1
            continue
        if line.strip().startswith("> "):
            add_para(doc, line.strip()[2:].strip(), size=10)
            i += 1
            continue

        # 表格
        if "|" in line and line.strip().startswith("|"):
            if _is_table_sep(line):
                i += 1
                continue
            table_buf.append(_parse_table_row(line))
            i += 1
            continue
        else:
            flush_table()

        stripped = line.strip()
        if not stripped:
            i += 1
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=1)
        elif stripped.startswith("- ") or stripped.startswith("• "):
            p = doc.add_paragraph(style="List Bullet")
            _add_rich_runs(p, stripped[2:].strip())
        elif re.match(r"^\d+\.\s", stripped):
            p = doc.add_paragraph(style="List Number")
            _add_rich_runs(p, re.sub(r"^\d+\.\s*", "", stripped))
        elif stripped.startswith("* ") and not stripped.startswith("**"):
            p = doc.add_paragraph(style="List Bullet")
            _add_rich_runs(p, stripped[2:].strip())
        else:
            p = doc.add_paragraph()
            _add_rich_runs(p, stripped)
        i += 1

    flush_table()


# ─── insights 逐词点评（Agent / 批量脚本共用）────────────────

def _metric_num(v, default=None):
    if v is None:
        return default
    if isinstance(v, float) and v != v:
        return default
    return v


def format_kw_metrics(kw: dict) -> str:
    parts = []
    if kw.get("level"):
        parts.append(str(kw["level"]))
    sv = _metric_num(kw.get("search_volume"))
    if sv is not None:
        parts.append(f"搜索量 {int(sv)}")
    bid = _metric_num(kw.get("bid"))
    if bid is not None:
        parts.append(f"竞价 ${bid}")
    conc = _metric_num(kw.get("concentration"))
    if conc is not None:
        parts.append(f"集中度 {conc}")
    sp = _metric_num(kw.get("sp_share"))
    if sp is not None:
        parts.append(f"SP份额 {sp}%")
    return "，".join(parts)


def comment_t1_keyword(kw: dict, stage: str, conc_median: float) -> str:
    """表1 TOP 词：竞争判断 + 阶段策略 + 匹配建议（禁止只复读数字）"""
    name = str(kw.get("keyword", ""))
    metrics = format_kw_metrics(kw)
    conc = _metric_num(kw.get("concentration"), conc_median or 0.3)
    match_type = str(kw.get("match_type") or "词组匹配")
    level = str(kw.get("level") or "")
    kl = name.lower()

    if conc_median and conc < conc_median * 0.85:
        comp = "集中度低于类目中位，TOP3 垄断不强"
    elif conc_median and conc > conc_median * 1.15:
        comp = "集中度高于类目中位，头部卖家占比较高"
    else:
        comp = "集中度接近类目中位，竞争密度适中"

    semantic = ""
    if any(x in kl for x in ("dog", "pet", "cat")):
        semantic = "宠物防抓场景，"
    elif "love seat" in kl or "loveseat" in kl:
        semantic = "双人位/loveseat 细分规格，"
    elif any(x in kl for x in ("washable", "waterproof", "stretch", "non slip")):
        semantic = "功能属性词，listing 需有对应卖点再投；"
    elif name in ("couch cover", "sofa cover", "sofa covers", "couch covers") or kl in ("couch cover", "sofa cover"):
        semantic = "泛品类大词、搜索意图较宽，"
    elif "for sofa" in kl or "for couch" in kl:
        semantic = "核心场景词（明确沙发套用途），"
    elif any(x in kl for x in ("fundas", "forros", "cubridores")):
        semantic = "西语词，建议单独低预算活动；"

    if stage == "新品期":
        if level == "S级":
            stage_tip = "新品期不宜 All-in 大词，先小预算验证转化；表1 PD 优先级多为「中」"
        else:
            stage_tip = "A 级适合词组试投，有订单再迁入精准组"
    elif stage == "成长期":
        stage_tip = "可逐步加码，优先有转化信号的规格词"
    else:
        stage_tip = "成熟期兼顾防守与补位，关注品牌与规格长尾"

    action = f"建议 **{match_type}**"
    if "精准" in match_type:
        action += "，单组控预算，7 天无转化暂停或否词"
    else:
        action += "，观察搜索词报告再收窄"

    return f"- **{name}**（{metrics}）：{semantic}{comp}。{stage_tip}。{action}。"


def render_t1_keyword_reviews(keywords: list, stage: str, conc_median: float, limit: int = 5) -> str:
    lines = [comment_t1_keyword(kw, stage, conc_median) for kw in keywords[:limit]]
    return "\n".join(lines) if lines else "- （无 S/A 级词）"


def comment_cross_keyword(kw: dict) -> str:
    name = str(kw.get("keyword", ""))
    grade = str(kw.get("grade") or kw.get("词级别") or "")
    metrics = format_kw_metrics({"level": grade, **kw})
    action = str(kw.get("action") or "")
    sources = str(kw.get("sources") or kw.get("出现来源") or "")
    sp = _metric_num(kw.get("sp_share"))
    tip = action or "按级别匹配方式试投"
    if sp is not None and sp < 5:
        tip += f"；竞品 SP 份额仅 {sp}%，抢位窗口尚可"
    elif sp is not None and sp >= 20:
        tip += f"；竞品 SP 份额 {sp}% 偏高，控预算防守"
    return f"- **{name}**（{metrics}，来源 {sources}）：{tip}。"


def render_cross_reviews(keywords: list, limit: int = 10) -> str:
    lines = [comment_cross_keyword(kw) for kw in keywords[:limit]]
    return "\n".join(lines) if lines else "- （无主攻词）"


# ─── LLM 写作提示（Agent 写 insights，不含图表）──────────────

def llm_prompt_template(report_key: str, stats: dict) -> str:
    """生成给 Agent 的写作提示：只写分析文字，图表由 Word 自动插入"""
    common_fmt = """
**格式要求**（供 Word 渲染，严格遵守）：
- 分节用 `## 标题`（不要用 `***` 或单 `*` 加粗）
- 列表用 `- ` 开头；加粗用 `**词**`
- 可用 `| 列 | 列 |` 表格，但复杂数据优先用列表
- **不要写图表节**，不要写 `![...](...)` — 图表由脚本自动插入 Word
- 不要写文档总标题（`# xxx`），Word 已有封面标题
"""

    templates = {
        "t1": common_fmt + """
**必须包含**：
1. 执行摘要（S/A/B/C 数量、竞争格局）
2. S/A 级投放策略（针对 {stage} 阶段）
3. **TOP5 重点词逐条点评** — 每条必须写：词性/场景、竞争（集中度 vs 类目中位）、{stage} 是否现在投、匹配方式与 listing 匹配度
4. PD 备战前 2 周行动清单

**TOP5 反例（禁止）**：「couch cover：couch cover，搜索量 88739，竞价 $1.31」——只复读数字不算点评
**TOP5 正例**：「**couch cover**（S级，搜索量 88739，竞价 $1.31，集中度 0.26）：泛品类大词，意图较宽；集中度低于类目中位。新品期词组匹配试投，7 天无转化则否词。」

Stats：
{stats_json}""",

        "t2": common_fmt + """
**必须包含**：
1. 竞品流量结构弱点总结
2. TOP10 高机会词抢位策略（列表）
3. PD 前广告活动搭建方案

Stats：
{stats_json}""",

        "t3": common_fmt + """
**必须包含**：
1. 缺口词战略价值（含 stats 中 defense_note：防守词=0 时必须解释原因与份额列模式）
2. TOP 缺口词逐条抢位建议（禁止只列数字）
3. PD 抢位行动清单

Stats：
{stats_json}""",

        "cross": common_fmt + """
**必须包含**：
1. 执行摘要：说明 SSS/SS/S 是「三表交叉分级」，**不是表1 S级**；S级=仅出现在一张筛选表中的词，数量多属正常
2. 交叉逻辑解读：三表各有多少词、两两交集各多少、为何 SSS 可能为 0（举 1~2 个具体词例）
3. **PD 主攻预算分配**（以 SSS+SS 为核心，S 级仅作扩展观察）
4. SSS 级词深度解读（若有）；SS 级 TOP10 分组投放方案
5. **PD 备战行动清单**（分周 checklist，含大促当天）

Stats：
{stats_json}""",

        "tracker": common_fmt + """
**必须包含**（词库动态对比，面向运营可执行）：
1. 执行摘要：上次 vs 本次词池变化，**重点说 SS/SSS 级**（S 级单表新增多为观察池膨胀，不必恐慌）
2. 新增词解读：按 SSS/SS/S 分级，每个 SS 级以上词写 1 句「要不要加预算」
3. 消失词解读：哪些从主攻池掉出，是否应降价/暂停
4. 搜索量/份额变动（若有）：对投放的影响
5. **本周投放调整行动清单**（逐条可执行，不要空泛模板）

对比窗口：{compare_note}
Stats：
{stats_json}""",
    }

    tpl = templates.get(report_key, "请根据以下数据写分析段：\n{stats_json}")
    return tpl.format(
        stats_json=json.dumps(stats, ensure_ascii=False, indent=2),
        stage=stats.get("stage", ""),
        product_category=stats.get("product_category", ""),
        compare_note=stats.get("compare_note", ""),
    )
