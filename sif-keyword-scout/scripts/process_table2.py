"""
process_table2.py - 处理表2：反查流量词
Usage: python process_table2.py --input "反查流量词_20260611.xlsx" --output-dir "处理结果/" --asin B0CRMP3RQT --date 20260611
"""
import argparse
import os
import sys
import warnings
warnings.filterwarnings("ignore")

from report_utils import (
    chart_path, get_col, stats_path, save_stats, load_insights, cleanup_intermediate_files,
    setup_doc_styles, add_para, apply_font_to_table, set_cell_text, finalize_doc_fonts,
    render_insights_section, add_doc_title, add_picture_centered, add_charts_section, insights_path,
)
from run_context import run_id as make_run_id

try:
    import pandas as pd
    import numpy as np
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.chart import BarChart, Reference, Series
except ImportError as e:
    print(f"缺少依赖：{e}\n请运行：pip install pandas openpyxl numpy", file=sys.stderr)
    sys.exit(1)

try:
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    plt.rcParams["font.family"] = ["Microsoft YaHei", "SimHei", "DejaVu Sans"]
    plt.rcParams["axes.unicode_minus"] = False
    HAS_MPL = True
except ImportError:
    HAS_MPL = False

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

THIN = Side(style="thin")
THIN_BORDER = Border(left=THIN, right=THIN, top=THIN, bottom=THIN)


def load_excel(path: str) -> pd.DataFrame:
    df = pd.read_excel(path, sheet_name=0, header=1)
    return df


def preprocess(df: pd.DataFrame) -> tuple:
    """定位关键列并计算基准统计"""
    col_kw    = get_col(df, ["关键词", "Keyword"])
    col_tag   = get_col(df, ["关键词效果类型", "效果类型", "Effect Type"])
    col_total = get_col(df, ["全部流量占比", "总流量占比"])
    col_nat   = get_col(df, ["自然流量占比", "Organic Traffic"])
    col_ad    = get_col(df, ["广告流量占比", "Ad Traffic"])
    col_conc = get_col(df, ["ABA TOP3集中度-点击", "ABA TOP3集中度-转化", "Top3集中度", "集中度"])
    if not col_conc and len(df.columns) >= 66:
        col_conc = df.columns[65]

    total_words = len(df)
    growth_count = 0
    if col_tag:
        growth_mask = df[col_tag].astype(str).str.contains("搜索量同比增长", na=False)
        growth_count = growth_mask.sum()
        if col_nat:
            nat_vals = pd.to_numeric(df.loc[growth_mask, col_nat], errors="coerce")
            weak_nat = (nat_vals < 0.1).sum()
        else:
            weak_nat = 0
    else:
        growth_mask = pd.Series([False] * len(df))
        weak_nat = 0

    print(f"\n📊 表2基准统计：")
    print(f"  总词数：{total_words}")
    print(f"  含「搜索量同比增长」标签词数：{growth_count}")
    print(f"  增长词中自然流量<10%（纯广告词）数：{weak_nat}")

    cols = dict(kw=col_kw, tag=col_tag, total=col_total, nat=col_nat, ad=col_ad, conc=col_conc)
    return df, cols, growth_count, weak_nat


def add_helper_columns(df: pd.DataFrame, cols: dict, params: dict | None = None) -> pd.DataFrame:
    df = df.copy()
    p = params or {}
    opp_conc = float(p.get("t2_opp_conc_max", 0.40))
    enter_low = float(p.get("t2_enter_low_conc", 0.25))
    enter_high = float(p.get("t2_enter_high_conc", 0.40))
    pure_ad_nat = float(p.get("t2_pure_ad_nat_max", 0.20))

    col_tag  = cols["tag"]
    col_nat  = cols["nat"]
    col_conc = cols["conc"]

    # 列1：是否增长词
    if col_tag:
        df["是否增长词"] = df[col_tag].astype(str).str.contains("搜索量同比增长", na=False).map({True: "是", False: "否"})
    else:
        df["是否增长词"] = "否"

    # 列2：竞品自然位状态
    def nat_status(val):
        try:
            v = float(val)
            if v > 0.6:
                return "稳定"
            elif v >= pure_ad_nat:
                return "未稳"
            else:
                return "纯广告"
        except (ValueError, TypeError):
            return "未知"

    if col_nat:
        df["竞品自然位状态"] = df[col_nat].apply(nat_status)
    else:
        df["竞品自然位状态"] = "未知"

    # 列3：抢位机会评级
    def opp_rating(row):
        is_growth = row.get("是否增长词", "否") == "是"
        nat_st    = row.get("竞品自然位状态", "稳定")
        try:
            conc = float(row.get(col_conc, 1) or 1) if col_conc else 1.0
        except (ValueError, TypeError):
            conc = 1.0

        if is_growth and nat_st in ("未稳", "纯广告") and conc < opp_conc:
            return "高"
        elif is_growth:
            return "中"
        else:
            return "低"

    df["抢位机会评级"] = df.apply(opp_rating, axis=1)

    # 列4：进入难度
    def enter_difficulty(row):
        nat_st = row.get("竞品自然位状态", "稳定")
        try:
            conc = float(row.get(col_conc, 0.5) or 0.5) if col_conc else 0.5
        except (ValueError, TypeError):
            conc = 0.5

        if conc < enter_low and nat_st in ("未稳", "纯广告"):
            return "低"
        elif conc > enter_high:
            return "高"
        else:
            return "中"

    df["进入难度"] = df.apply(enter_difficulty, axis=1)

    # 列5：建议操作
    def suggest_op(row):
        rating  = row.get("抢位机会评级", "低")
        diff    = row.get("进入难度", "中")
        if rating == "高" and diff == "低":
            return "PD前开独立广告活动，精准匹配抢位"
        elif rating == "高" and diff == "中":
            return "词组匹配测试，控预算看转化再加码"
        elif rating == "中":
            return "广泛匹配跟打，不单独建组"
        else:
            return "暂不投入，列入观察"

    df["建议操作"] = df.apply(suggest_op, axis=1)
    return df


def build_board_sheet(wb: openpyxl.Workbook, df: pd.DataFrame, cols: dict):
    if "竞品弱点看板" in wb.sheetnames:
        del wb["竞品弱点看板"]
    ws = wb.create_sheet("竞品弱点看板")
    ws.sheet_view.showGridLines = False

    head_fill = PatternFill(start_color="843C0C", end_color="843C0C", fill_type="solid")
    head_font = Font(color="FFFFFF", bold=True, name="Microsoft YaHei", size=10)
    title_font = Font(bold=True, name="Microsoft YaHei", size=13)
    body_font  = Font(name="Microsoft YaHei", size=10)

    row = 1

    # ── 统计1：增长词vs非增长词 ──
    ws.cell(row, 1, "📊 增长词 vs 非增长词").font = title_font
    row += 1
    for ci, h in enumerate(["类型", "词数量", "占比"], 1):
        c = ws.cell(row, ci, h); c.fill = head_fill; c.font = head_font; c.border = THIN_BORDER
    row += 1
    total = len(df)
    growth_cnt = (df["是否增长词"] == "是").sum()
    non_cnt = total - growth_cnt
    for label, cnt in [("增长词", growth_cnt), ("非增长词", non_cnt)]:
        ws.cell(row, 1, label).border = THIN_BORDER
        ws.cell(row, 2, cnt).border = THIN_BORDER
        ws.cell(row, 3, f"{cnt/total*100:.1f}%" if total else "0%").border = THIN_BORDER
        row += 1
    row += 1

    # ── 统计2：机会评级 ──
    ws.cell(row, 1, "📊 抢位机会评级分布").font = title_font
    row += 1
    for ci, h in enumerate(["评级", "词数量", "占比"], 1):
        c = ws.cell(row, ci, h); c.fill = head_fill; c.font = head_font; c.border = THIN_BORDER
    row += 1
    rating_start = row
    for rating in ["高", "中", "低"]:
        cnt = (df["抢位机会评级"] == rating).sum()
        ws.cell(row, 1, rating).border = THIN_BORDER
        ws.cell(row, 2, cnt).border = THIN_BORDER
        ws.cell(row, 3, f"{cnt/total*100:.1f}%" if total else "0%").border = THIN_BORDER
        row += 1
    rating_end = row - 1
    row += 1

    # ── TOP15机会词 ──
    ws.cell(row, 1, "🎯 TOP15机会词清单（高机会优先 + 全部流量占比降序）").font = title_font
    row += 1
    col_kw    = cols["kw"]
    col_total = cols["total"]
    col_nat   = cols["nat"]
    col_ad    = cols["ad"]

    board_headers = ["关键词", "全部流量占比%", "自然流量占比%", "广告流量占比%", "竞品自然位状态", "进入难度", "建议操作"]
    for ci, h in enumerate(board_headers, 1):
        c = ws.cell(row, ci, h); c.fill = head_fill; c.font = head_font; c.border = THIN_BORDER
    row += 1

    top_df = df[df["抢位机会评级"] == "高"].copy()
    if col_total:
        top_df[col_total] = pd.to_numeric(top_df[col_total], errors="coerce")
        top_df = top_df.sort_values(col_total, ascending=False)
    top_df = top_df.head(15)

    for i, (_, r) in enumerate(top_df.iterrows()):
        fill = PatternFill(start_color="FFF2CC", end_color="FFF2CC", fill_type="solid") if i % 2 == 0 else None
        def fmt_pct(val):
            try: return f"{float(val)*100:.1f}%"
            except: return ""
        vals = [
            str(r.get(col_kw, "")) if col_kw else "",
            fmt_pct(r.get(col_total)) if col_total else "",
            fmt_pct(r.get(col_nat))   if col_nat  else "",
            fmt_pct(r.get(col_ad))    if col_ad   else "",
            str(r.get("竞品自然位状态", "")),
            str(r.get("进入难度", "")),
            str(r.get("建议操作", "")),
        ]
        for ci, v in enumerate(vals, 1):
            c = ws.cell(row, ci, v)
            c.border = THIN_BORDER
            c.font   = Font(name="Microsoft YaHei", size=9)
            if fill: c.fill = fill
        row += 1
    row += 1

    # ── 嵌入图表 ──
    # 图1：高/中/低评级数量（用统计数据行）
    chart_data_row = row
    for ri, rating in enumerate(["高", "中", "低"]):
        cnt = (df["抢位机会评级"] == rating).sum()
        ws.cell(chart_data_row + ri, 10, rating)
        ws.cell(chart_data_row + ri, 11, cnt)

    bc1 = BarChart()
    bc1.type = "col"; bc1.title = "抢位机会评级分布"
    bc1.y_axis.title = "词数量"; bc1.width = 14; bc1.height = 9
    data_ref = Reference(ws, min_col=11, min_row=chart_data_row, max_row=chart_data_row + 2)
    cats_ref = Reference(ws, min_col=10, min_row=chart_data_row, max_row=chart_data_row + 2)
    s1 = Series(data_ref, title="词数量"); bc1.series.append(s1); bc1.set_categories(cats_ref)
    ws.add_chart(bc1, f"A{row}")

    # 图二：TOP8机会词 全部流量 vs 自然流量 并排柱
    top8 = top_df.head(8).copy()
    cmp_row = row + 1
    ws.cell(row, 13, "关键词")
    ws.cell(row, 14, "全部流量")
    ws.cell(row, 15, "自然流量")
    for i, (_, r) in enumerate(top8.iterrows()):
        kw = str(r.get(col_kw, ""))[:8]
        total_v = float(r.get(col_total, 0) or 0) * 100 if col_total else 0
        nat_v = float(r.get(col_nat, 0) or 0) * 100 if col_nat else 0
        ws.cell(cmp_row + i, 13, kw)
        ws.cell(cmp_row + i, 14, round(total_v, 1))
        ws.cell(cmp_row + i, 15, round(nat_v, 1))
    if len(top8) > 0:
        bc2 = BarChart()
        bc2.type = "col"
        bc2.grouping = "clustered"
        bc2.title = "TOP8机会词：全部流量 vs 自然流量（橙柱低=广告依赖高）"
        bc2.y_axis.title = "流量占比(%)"
        bc2.width = 18
        bc2.height = 10
        data2 = Reference(ws, min_col=14, min_row=row, max_col=15, max_row=cmp_row + len(top8) - 1)
        cats2 = Reference(ws, min_col=13, min_row=cmp_row, max_row=cmp_row + len(top8) - 1)
        bc2.add_data(data2, titles_from_data=True)
        bc2.set_categories(cats2)
        ws.add_chart(bc2, f"J{chart_data_row}")

    row += 18

    # 调整列宽
    for ci, w in enumerate([22, 16, 16, 16, 16, 12, 30], 1):
        ws.column_dimensions[openpyxl.utils.get_column_letter(ci)].width = w


def write_helper_cols_to_main(wb: openpyxl.Workbook, df: pd.DataFrame):
    ws = wb.active
    if not ws:
        return
    max_col = ws.max_column
    new_cols = ["是否增长词", "竞品自然位状态", "抢位机会评级", "进入难度", "建议操作"]
    head_fill = PatternFill(start_color="843C0C", end_color="843C0C", fill_type="solid")
    head_font = Font(color="FFFFFF", bold=True, name="Microsoft YaHei")

    for ci, hdr in enumerate(new_cols, max_col + 1):
        c = ws.cell(2, ci, hdr)
        c.fill = head_fill; c.font = head_font; c.border = THIN_BORDER

    for ri, (_, row) in enumerate(df.iterrows()):
        excel_row = ri + 3
        for ci, col in enumerate(new_cols, max_col + 1):
            c = ws.cell(excel_row, ci, row.get(col, ""))
            c.border = THIN_BORDER
            c.font = Font(name="Microsoft YaHei", size=9)
            if col == "抢位机会评级":
                color_map = {"高": "C6EFCE", "中": "FFEB9C", "低": "FCE4D6"}
                fc = color_map.get(str(row.get(col, "")), "FFFFFF")
                c.fill = PatternFill(start_color=fc, end_color=fc, fill_type="solid")


def build_t2_stats(df, cols, asin, date) -> dict:
    col_kw = cols["kw"]
    col_total = cols["total"]
    top_df = df[df["抢位机会评级"] == "高"].copy()
    if col_total:
        top_df[col_total] = pd.to_numeric(top_df[col_total], errors="coerce")
        top_df = top_df.sort_values(col_total, ascending=False)
    top_list = []
    for _, r in top_df.head(15).iterrows():
        top_list.append({
            "keyword": str(r.get(col_kw, "")) if col_kw else "",
            "total_traffic_pct": float(r.get(col_total, 0) or 0) if col_total else None,
            "nat_status": str(r.get("竞品自然位状态", "")),
            "difficulty": str(r.get("进入难度", "")),
            "action": str(r.get("建议操作", "")),
        })
    return {
        "report": "t2_competitor_weakness",
        "asin": asin, "date": date,
        "total": len(df),
        "growth_count": int((df["是否增长词"] == "是").sum()),
        "high_opp": int((df["抢位机会评级"] == "高").sum()),
        "pure_ad_count": int((df["竞品自然位状态"] == "纯广告").sum()),
        "top_opportunity_keywords": top_list,
    }


def build_word_report(df, cols, asin, date, output_dir, insights=""):
    if not HAS_DOCX:
        print("⚠️ 未安装 python-docx，跳过 Word", file=sys.stderr)
        return None

    doc = Document()
    setup_doc_styles(doc)
    add_doc_title(doc, f"{asin} 竞品弱点分析报告", f"生成日期：{date}")

    total = len(df)
    growth_cnt = (df["是否增长词"] == "是").sum()
    high_cnt = (df["抢位机会评级"] == "高").sum()
    mid_cnt = (df["抢位机会评级"] == "中").sum()
    pure_ad_cnt = (df["竞品自然位状态"] == "纯广告").sum()

    col_kw = cols["kw"]
    col_total = cols["total"]
    col_nat = cols["nat"]

    top_df = df[df["抢位机会评级"] == "高"].copy()
    if col_total:
        top_df[col_total] = pd.to_numeric(top_df[col_total], errors="coerce")
        top_df = top_df.sort_values(col_total, ascending=False)
    top5_kw = "、".join(top_df.head(5)[col_kw].astype(str).tolist()) if col_kw and len(top_df) > 0 else "（无）"

    doc.add_heading("一、执行摘要", level=1)
    add_para(doc,
        f"共分析 {total} 个关键词。增长词 {growth_cnt} 个；高机会词 {high_cnt} 个，中机会词 {mid_cnt} 个；"
        f"纯广告词 {pure_ad_cnt} 个。高机会 TOP5：{top5_kw}。"
    )

    doc.add_heading("二、策略洞察与抢位建议", level=1)
    render_insights_section(doc, insights, heading="")

    doc.add_heading("三、TOP15机会词清单", level=1)
    top_df = top_df.head(15)
    if len(top_df) > 0:
        tbl = doc.add_table(rows=len(top_df) + 1, cols=5)
        tbl.style = "Table Grid"
        for ci, h in enumerate(["关键词", "全部流量%", "自然流量%", "竞品自然位", "建议操作"]):
            set_cell_text(tbl.rows[0].cells[ci], h)
        for ri, (_, r) in enumerate(top_df.iterrows(), 1):
            def fmt(v):
                try:
                    return f"{float(v) * 100:.1f}%"
                except Exception:
                    return ""
            set_cell_text(tbl.rows[ri].cells[0], str(r.get(col_kw, "")) if col_kw else "")
            set_cell_text(tbl.rows[ri].cells[1], fmt(r.get(col_total)) if col_total else "")
            set_cell_text(tbl.rows[ri].cells[2], fmt(r.get(col_nat)) if col_nat else "")
            set_cell_text(tbl.rows[ri].cells[3], str(r.get("竞品自然位状态", "")))
            set_cell_text(tbl.rows[ri].cells[4], str(r.get("建议操作", "")))
        apply_font_to_table(tbl)

    if HAS_MPL:
        chart_paths = _gen_charts_t2(df, cols, output_dir, asin, date)
        items = []
        if len(chart_paths) >= 1:
            items.append(("图1 · 饼图：抢位机会评级分布（高/中/低）", chart_paths[0]))
        if len(chart_paths) >= 2:
            items.append(("图2 · 分组柱形图：TOP 高机会词 全部流量占比 vs 自然流量占比", chart_paths[1]))
        add_charts_section(doc, items, section_title="四、可视化图表", level=1)

    finalize_doc_fonts(doc)
    out = os.path.join(output_dir, f"{asin}_竞品弱点分析报告_{date}.docx")
    doc.save(out)
    print(f"✅ Word报告已生成：{out}")
    return out


def _gen_charts_t2(df, cols, output_dir, asin, date):
    paths = []
    try:
        labels = ["高", "中", "低"]
        counts = [(df["抢位机会评级"] == r).sum() for r in labels]
        colors = ["#FF6B35", "#FFC000", "#A9A9A9"]
        fig, ax = plt.subplots(figsize=(7, 5))
        non_zero = [(l, c, col) for l, c, col in zip(labels, counts, colors) if c > 0]
        if non_zero:
            pie_labels, pie_counts, pie_colors = zip(*non_zero)
            wedges, texts, autotexts = ax.pie(
                pie_counts, labels=pie_labels, colors=pie_colors, autopct="%1.0f%%",
                startangle=90, textprops={"fontsize": 11},
            )
            for t in autotexts:
                t.set_fontweight("bold")
        ax.set_title("抢位机会评级分布", fontsize=13, fontweight="bold")
        p1 = chart_path(output_dir, f"chart1_t2_{asin}_{date}.png")
        fig.savefig(p1, dpi=150, bbox_inches="tight"); plt.close(fig); paths.append(p1)
    except Exception as e:
        print(f"图1失败: {e}", file=sys.stderr)

    try:
        col_kw    = cols["kw"]
        col_total = cols["total"]
        col_nat   = cols["nat"]
        if col_total and col_nat:
            top_df = df[df["抢位机会评级"] == "高"].copy()
            top_df[col_total] = pd.to_numeric(top_df[col_total], errors="coerce")
            top_df[col_nat]   = pd.to_numeric(top_df[col_nat],   errors="coerce")
            top_df = top_df.sort_values(col_total, ascending=False).head(15)
            if len(top_df) > 0:
                labels = top_df[col_kw].astype(str).str[:10].tolist() if col_kw else list(range(len(top_df)))
                x = np.arange(len(labels)); width = 0.35
                fig, ax = plt.subplots(figsize=(max(10, len(labels)), 5))
                ax.bar(x - width/2, top_df[col_total] * 100, width, label="全部流量%", color="#4472C4")
                ax.bar(x + width/2, top_df[col_nat]   * 100, width, label="自然流量%", color="#ED7D31")
                ax.set_xticks(x); ax.set_xticklabels(labels, rotation=45, ha="right", fontsize=8)
                ax.set_title("TOP15机会词 全部流量 vs 自然流量", fontsize=12, fontweight="bold")
                ax.set_ylabel("%")
                ax.legend()
                ax.spines[["top", "right"]].set_visible(False)
                p2 = chart_path(output_dir, f"chart2_t2_{asin}_{date}.png")
                fig.savefig(p2, dpi=150, bbox_inches="tight"); plt.close(fig); paths.append(p2)
    except Exception as e:
        print(f"图2失败: {e}", file=sys.stderr)
    return paths


def main():
    from datetime import date as _date, datetime
    from threshold_presets import add_product_type_arg, resolve_thresholds, format_params_summary, collect_overrides_from_args

    parser = argparse.ArgumentParser(description="处理Sif反查流量词（表2）")
    parser.add_argument("--input",       required=True)
    parser.add_argument("--output-dir",  required=True)
    parser.add_argument("--asin",        required=True)
    parser.add_argument("--date",        default=_date.today().strftime("%Y%m%d"),
                        help="日期YYYYMMDD，默认今天")
    parser.add_argument("--stage",       default="新品期", choices=["新品期", "成长期", "成熟期"])
    add_product_type_arg(parser)
    parser.add_argument("--t2-opp-conc-max", type=float, default=None, help="高机会词集中度上限")
    parser.add_argument("--t2-pure-ad-nat-max", type=float, default=None, help="纯广告判定：自然流量低于此比例")
    parser.add_argument("--insights",      default="")
    parser.add_argument("--insights-file", default="")
    parser.add_argument("--skip-word",     action="store_true")
    args = parser.parse_args()
    args.insights = load_insights(args.insights, args.insights_file)
    params = resolve_thresholds(args.product_type, args.stage, collect_overrides_from_args(args))

    if not os.path.exists(args.input):
        print(f"❌ 找不到文件：{args.input}", file=sys.stderr); sys.exit(1)

    rid = make_run_id(args.date)
    os.makedirs(args.output_dir, exist_ok=True)
    print(f"📂 读取文件：{args.input}")
    print(f"  预设：{format_params_summary(params)}")
    df = load_excel(args.input)
    print(f"📊 共读取 {len(df)} 行")

    df, cols, growth_cnt, weak_nat = preprocess(df)
    df = add_helper_columns(df, cols, params)

    print(f"\n  高机会词：{(df['抢位机会评级']=='高').sum()} | 中：{(df['抢位机会评级']=='中').sum()} | 低：{(df['抢位机会评级']=='低').sum()}")

    import shutil
    out_excel = os.path.join(args.output_dir, f"{args.asin}_竞品弱点分析_{rid}.xlsx")
    shutil.copy2(args.input, out_excel)

    print("📝 写入辅助列...")
    wb = openpyxl.load_workbook(out_excel)
    write_helper_cols_to_main(wb, df)
    build_board_sheet(wb, df, cols)
    wb.save(out_excel)
    print(f"✅ Excel已生成：{out_excel}")

    stats = build_t2_stats(df, cols, args.asin, args.date)
    stats["run_id"] = rid
    sp = save_stats(stats, stats_path(args.output_dir, args.asin, "t2", rid))
    print(f"📊 Stats JSON：{sp}")

    if args.skip_word:
        print("\n" + "=" * 60)
        print("📝 [Agent] 写 insights_t2.md 后带 --insights-file 重建 Word")
        print(f"  stats：{sp}")
        print(f"  insights：{insights_path(args.output_dir, 't2')}")
        print("=" * 60)
    else:
        build_word_report(df, cols, args.asin, rid, args.output_dir, insights=args.insights)
        cleanup_intermediate_files(args.output_dir, args.asin, rid, "t2")

    print(f"\n✅ 表2处理完成！输出目录：{args.output_dir}")


if __name__ == "__main__":
    main()
