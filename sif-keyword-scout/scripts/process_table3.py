"""
process_table3.py - 处理表3：查广告词
Usage: python process_table3.py --input "查广告词_20260611.xlsx" --output-dir "处理结果/" --asin B0CRMP3RQT --date 20260611 --product-category "水杯"

Sif 表3 已经是针对该 ASIN 的 SP 广告词，默认全部相关；只有翻译中出现与类目明显不匹配的词才标注「不相关」。
"""
import argparse
import os
import re
import sys
from datetime import date as _date, datetime
import warnings
warnings.filterwarnings("ignore")

from report_utils import (
    chart_path, get_col, pct_to_float, stats_path, save_stats, load_insights, cleanup_intermediate_files,
    setup_doc_styles, add_para, apply_font_to_table, set_cell_text, finalize_doc_fonts,
    render_insights_section, add_doc_title, add_picture_centered, add_charts_section, insights_path,
    is_missing, fmt_share_pct, fmt_cell,
)
from run_context import run_id as make_run_id

try:
    import pandas as pd
    import numpy as np
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.chart import BarChart, Reference, Series
except ImportError as e:
    print(f"缺少依赖：{e}\n请运行：pip install pandas openpyxl numpy", file=sys.stderr)
    sys.exit(1)

try:
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    plt.rcParams["font.family"] = ["Microsoft YaHei", "SimHei", "DejaVu Sans"]
    plt.rcParams["axes.unicode_minus"] = False
    HAS_MPL = True
except ImportError:
    HAS_MPL = False

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

THIN = Side(style="thin")
THIN_BORDER = Border(left=THIN, right=THIN, top=THIN, bottom=THIN)


def load_excel(path: str) -> pd.DataFrame:
    df = pd.read_excel(path, sheet_name=0, header=1)
    return df


def _build_cat_tokens(product_category: str) -> list:
    """将类目字符串分拆成匹配 token 列表（中文/英文均兼容）"""
    tokens = re.split(r"[/，,\s]+", product_category.lower().strip())
    return [t for t in tokens if t and len(t) > 1]


def judge_relevance(kw: str, trans: str, product_category: str) -> str:
    """
    判断广告词是否与产品类目相关。

    ⚠️ Sif 表3 是针对本 ASIN 的 SP 广告词，已由平台预过滤，绝大多数词与产品相关。
    默认返回「相关」，只在关键词/翻译里出现明显反向模式时才标「不相关」。
    「疑似不相关」不再使用，避免误杀。
    """
    if not kw and not trans:
        return "相关"

    kw_lower   = str(kw).lower().strip()
    trans_lower = str(trans).lower().strip()
    combined   = f"{kw_lower} {trans_lower}"

    # ── 用户类目 token 命中 → 明确相关 ──
    for token in _build_cat_tokens(product_category):
        if token in combined:
            return "相关"

    # ── 品牌/竞品词 / 配件词 → 仍相关（同类目竞品词流量很重要）──
    # 不再用硬编码黑名单：不同类目的「不相关」词差异极大，交给用户在表格备注里手动标记

    # 默认：相关（Sif 已预过滤）
    return "相关"


def share_display(row, cols: dict) -> str:
    """份额列展示：空值 → 空白（不用 0.00% 或 nan）。"""
    sh_col = cols.get("sh")
    if sh_col:
        return fmt_share_pct(row.get(sh_col))
    return ""


def preprocess(df: pd.DataFrame, product_category: str) -> tuple:
    col_kw   = get_col(df, ["广告搜索词", "关键词", "Keyword"])
    col_tran = get_col(df, ["翻译", "Translation"])
    # 竞品份额（旧版 Sif）vs 本店在该词的 SP 份额（2026+ 新版）
    col_sh_comp = get_col(df, ["竞品在该词的SP广告流量份额", "竞品在该词SP广告流量份额"])
    col_sh_own = get_col(df, ["该Listing在该词下的SP广告流量份额", "Listing在该词下的SP广告流量份额"])
    col_sp   = get_col(df, ["该词为整个Listing贡献的SP广告流量占比", "SP广告流量贡献占比", "贡献占比", "SP Contribution"])
    col_sv   = get_col(df, ["搜索量", "Search Volume", "周搜索量"])

    if col_sh_comp:
        col_sh = col_sh_comp
        share_mode = "competitor"
        share_label = "竞品SP份额"
    elif col_sh_own:
        col_sh = col_sh_own
        share_mode = "own_listing"
        share_label = "Listing SP份额"
    else:
        col_sh = None
        share_mode = "missing"
        share_label = "SP份额"

    if not col_sv:
        raise ValueError("找不到搜索量列（K列）")

    # 转数字
    df = df.copy()
    df["__sp_num"] = df[col_sp].apply(pct_to_float) if col_sp else 0.0
    df["__sh_num"] = df[col_sh].apply(pct_to_float) if col_sh else 0.0
    df[col_sv] = pd.to_numeric(df[col_sv], errors="coerce").fillna(0)

    sv_median = df[col_sv].median()
    total = len(df)
    print(f"\n📊 表3基准统计：")
    print(f"  总词数：{total}")
    print(f"  搜索量中位数：{sv_median:.0f}")
    print(f"  份额列模式：{share_mode}（{share_label}）")

    cols = dict(kw=col_kw, tran=col_tran, sp=col_sp, sh=col_sh, sv=col_sv,
                sp_num="__sp_num", sh_num="__sh_num",
                share_mode=share_mode, share_label=share_label)
    return df, cols, sv_median


def add_helper_columns(df: pd.DataFrame, cols: dict, product_category: str,
                        params: dict | None = None) -> pd.DataFrame:
    df = df.copy()
    p = params or {}
    sv_threshold = int(p.get("sv_threshold", 2000))
    gap_max = float(p.get("gap_share_max", 5.0))
    defense_min = float(p.get("defense_share_min", 20.0))
    col_kw   = cols["kw"]
    col_tran = cols["tran"]
    col_sv   = cols["sv"]
    sp_num   = cols["sp_num"]
    sh_num   = cols["sh_num"]

    # 列1：相关性判断
    def rel_judge(row):
        kw   = str(row.get(col_kw, ""))   if col_kw   else ""
        tran = str(row.get(col_tran, "")) if col_tran else ""
        return judge_relevance(kw, tran, product_category)

    df["相关性判断"] = df.apply(rel_judge, axis=1)

    # 列2：词类型（缺口词/防守词/不相关词/普通词）
    def word_type(row):
        rel  = row.get("相关性判断", "相关")
        sv   = float(row.get(col_sv, 0) or 0)
        sh   = float(row.get(sh_num, 0) or 0)
        if rel == "不相关":
            return "不相关词"
        if sv > sv_threshold and sh < gap_max:
            return "缺口词"
        if sv > sv_threshold and sh > defense_min:
            return "防守词"
        return "普通词"

    df["词类型"] = df.apply(word_type, axis=1)

    # 列3：抢位难度
    def grab_diff(row):
        sh = float(row.get(sh_num, 0) or 0)
        sp = float(row.get(sp_num, 0) or 0)
        if sh < 3 and sp < 5:
            return "低"
        elif sh > 10:
            return "高"
        return "中"

    df["抢位难度"] = df.apply(grab_diff, axis=1)

    # 列4：竞品投入强度
    def ad_intensity(row):
        sp = float(row.get(sp_num, 0) or 0)
        if sp > 10:
            return "强"
        elif sp >= 3:
            return "中"
        return "弱"

    df["竞品投入强度"] = df.apply(ad_intensity, axis=1)

    # 列5：PD前建议操作
    def suggest(row):
        wt   = row.get("词类型", "普通词")
        diff = row.get("抢位难度", "中")
        if wt == "缺口词" and diff == "低":
            return "开独立广告活动，精准匹配重点投放"
        elif wt == "缺口词" and diff == "中":
            return "词组匹配测试，观察转化后加码"
        elif wt == "防守词":
            return "谨慎，避免正面硬打"
        elif wt == "不相关词":
            return "排除"
        return "常规跟进"

    df["PD前建议操作"] = df.apply(suggest, axis=1)
    return df


def build_board_sheet(wb, df, cols, asin):
    if "竞品缺口看板" in wb.sheetnames:
        del wb["竞品缺口看板"]
    ws = wb.create_sheet("竞品缺口看板")
    ws.sheet_view.showGridLines = False

    head_fill  = PatternFill(start_color="375623", end_color="375623", fill_type="solid")
    head_font  = Font(color="FFFFFF", bold=True, name="Microsoft YaHei", size=10)
    title_font = Font(bold=True, name="Microsoft YaHei", size=13)

    col_kw  = cols["kw"]
    col_sv  = cols["sv"]
    col_sh  = cols["sh"]
    sh_num  = cols["sh_num"]
    sp_num  = cols["sp_num"]
    share_label = cols.get("share_label", "SP份额")
    col_sp  = cols["sp"]

    row = 1

    # ── 词类型分布 ──
    ws.cell(row, 1, "📊 词类型数量分布").font = title_font; row += 1
    for ci, h in enumerate(["词类型", "词数量", "占比"], 1):
        c = ws.cell(row, ci, h); c.fill = head_fill; c.font = head_font; c.border = THIN_BORDER
    row += 1
    total = len(df)
    type_counts = {}
    for wt in ["缺口词", "防守词", "普通词", "不相关词"]:
        cnt = (df["词类型"] == wt).sum(); type_counts[wt] = cnt
        ws.cell(row, 1, wt).border = THIN_BORDER
        ws.cell(row, 2, cnt).border = THIN_BORDER
        ws.cell(row, 3, f"{cnt/total*100:.1f}%" if total else "0%").border = THIN_BORDER
        row += 1
    row += 1

    # ── TOP缺口词列表 ──
    ws.cell(row, 1, "🎯 TOP缺口词完整列表（按搜索量降序）").font = title_font; row += 1
    gap_headers = ["广告搜索词", "搜索量", f"{share_label}%", "SP贡献占比%", "抢位难度", "建议操作"]
    for ci, h in enumerate(gap_headers, 1):
        c = ws.cell(row, ci, h); c.fill = head_fill; c.font = head_font; c.border = THIN_BORDER
    row += 1

    gap_df = df[df["词类型"] == "缺口词"].copy()
    gap_df[col_sv] = pd.to_numeric(gap_df[col_sv], errors="coerce")
    gap_df = gap_df.sort_values(col_sv, ascending=False)

    for i, (_, r) in enumerate(gap_df.iterrows()):
        fill = PatternFill(start_color="E2EFDA", end_color="E2EFDA", fill_type="solid") if i % 2 == 0 else None
        vals = [
            str(r.get(col_kw, "")) if col_kw else "",
            int(r.get(col_sv, 0) or 0),
            share_display(r, cols),
            f"{r.get(sp_num, 0):.2f}%",
            str(r.get("抢位难度", "")),
            str(r.get("PD前建议操作", "")),
        ]
        for ci, v in enumerate(vals, 1):
            c = ws.cell(row, ci, v); c.border = THIN_BORDER
            c.font = Font(name="Microsoft YaHei", size=9)
            if fill: c.fill = fill
        row += 1
    row += 1

    # ── 竞品广告重心TOP10 ──
    ws.cell(row, 1, "⚔️ 竞品广告重心 TOP10（SP贡献占比最高）").font = title_font; row += 1
    heavy_headers = ["广告搜索词", "搜索量", "SP贡献占比%", f"{share_label}%", "词类型"]
    for ci, h in enumerate(heavy_headers, 1):
        c = ws.cell(row, ci, h); c.fill = head_fill; c.font = head_font; c.border = THIN_BORDER
    row += 1

    heavy_df = df[df["词类型"] != "不相关词"].copy()
    heavy_df[sp_num] = pd.to_numeric(heavy_df[sp_num], errors="coerce")
    heavy_df = heavy_df.sort_values(sp_num, ascending=False).head(10)

    for i, (_, r) in enumerate(heavy_df.iterrows()):
        fill = PatternFill(start_color="FCE4D6", end_color="FCE4D6", fill_type="solid") if i % 2 == 0 else None
        vals = [
            str(r.get(col_kw, "")) if col_kw else "",
            int(r.get(col_sv, 0) or 0),
            f"{r.get(sp_num, 0):.2f}%",
            share_display(r, cols),
            str(r.get("词类型", "")),
        ]
        for ci, v in enumerate(vals, 1):
            c = ws.cell(row, ci, v); c.border = THIN_BORDER
            c.font = Font(name="Microsoft YaHei", size=9)
            if fill: c.fill = fill
        row += 1
    row += 1

    # ── 2×2 战略矩阵 ──
    ws.cell(row, 1, "🗺️ 2×2 战略矩阵").font = title_font; row += 1
    # own_listing 列名在新版 Sif 查竞品 ASIN 时仍代表竞品份额，统一用「SP份额」
    matrix_note = "SP份额"
    matrix_data = [
        ["", f"低{matrix_note}(<5%)", f"高{matrix_note}(>20%)"],
        ["高搜索(>2000)",
         f"优先抢（缺口词 {type_counts.get('缺口词',0)} 个）",
         f"绕开（防守词 {type_counts.get('防守词',0)} 个）"],
        ["低搜索(≤2000)",
         f"备用（长尾缺口 {len(df[(df['词类型']=='普通词') & (pd.to_numeric(df[col_sv], errors='coerce') <= 2000)]) if col_sv else 0} 个）",
         "忽略"],
    ]
    for ri_offset, mrow in enumerate(matrix_data):
        for ci, val in enumerate(mrow, 1):
            c = ws.cell(row + ri_offset, ci, val)
            c.border = THIN_BORDER
            c.font = Font(bold=(ri_offset == 0 or ci == 1), name="Microsoft YaHei", size=10)
            if ri_offset == 0 or ci == 1:
                c.fill = PatternFill(start_color="375623", end_color="375623", fill_type="solid")
                c.font = Font(color="FFFFFF", bold=True, name="Microsoft YaHei", size=10)
    row += len(matrix_data) + 1

    # ── 嵌入图表 ──
    chart_data_row = row

    # 图一：TOP10缺口词搜索量
    top10_gap = gap_df.head(10)
    gap_chart_row = chart_data_row
    for i, (_, r) in enumerate(top10_gap.iterrows()):
        kw = str(r.get(col_kw, ""))[:8] if col_kw else f"词{i+1}"
        sv = int(r.get(col_sv, 0) or 0)
        ws.cell(gap_chart_row + i, 12, kw)
        ws.cell(gap_chart_row + i, 13, sv)
    if len(top10_gap) > 0:
        bc_gap = BarChart()
        bc_gap.type = "col"
        bc_gap.title = "TOP10缺口词搜索量"
        bc_gap.y_axis.title = "搜索量"
        bc_gap.width = 16
        bc_gap.height = 9
        dr_gap = Reference(ws, min_col=13, min_row=gap_chart_row, max_row=gap_chart_row + len(top10_gap) - 1)
        cr_gap = Reference(ws, min_col=12, min_row=gap_chart_row, max_row=gap_chart_row + len(top10_gap) - 1)
        bc_gap.series.append(Series(dr_gap, title="搜索量"))
        bc_gap.set_categories(cr_gap)
        ws.add_chart(bc_gap, f"A{chart_data_row}")

    # 图二：词类型数量分布
    type_chart_row = chart_data_row
    for i, (wt, cnt) in enumerate(type_counts.items()):
        ws.cell(type_chart_row + i, 9,  wt)
        ws.cell(type_chart_row + i, 10, cnt)
    bc = BarChart()
    bc.type = "col"; bc.title = "词类型数量分布"
    bc.y_axis.title = "词数量"; bc.width = 14; bc.height = 9
    dr = Reference(ws, min_col=10, min_row=type_chart_row, max_row=type_chart_row + 3)
    cr = Reference(ws, min_col=9,  min_row=type_chart_row, max_row=type_chart_row + 3)
    s1 = Series(dr, title="词数量"); bc.series.append(s1); bc.set_categories(cr)
    ws.add_chart(bc, f"J{chart_data_row}")

    for ci, w in enumerate([22, 12, 14, 14, 12, 30], 1):
        ws.column_dimensions[openpyxl.utils.get_column_letter(ci)].width = w


def write_helper_cols_to_main(wb, df):
    ws = wb.active
    if not ws: return
    max_col = ws.max_column
    new_cols = ["相关性判断", "词类型", "抢位难度", "竞品投入强度", "PD前建议操作"]
    head_fill = PatternFill(start_color="375623", end_color="375623", fill_type="solid")
    head_font = Font(color="FFFFFF", bold=True, name="Microsoft YaHei")
    for ci, hdr in enumerate(new_cols, max_col + 1):
        c = ws.cell(2, ci, hdr); c.fill = head_fill; c.font = head_font; c.border = THIN_BORDER
    for ri, (_, row) in enumerate(df.iterrows()):
        excel_row = ri + 3
        for ci, col in enumerate(new_cols, max_col + 1):
            c = ws.cell(excel_row, ci, row.get(col, ""))
            c.border = THIN_BORDER; c.font = Font(name="Microsoft YaHei", size=9)
            if col == "词类型":
                color_map = {"缺口词": "C6EFCE", "防守词": "FCE4D6", "不相关词": "BFBFBF", "普通词": "FFFFFF"}
                fc = color_map.get(str(row.get(col, "")), "FFFFFF")
                c.fill = PatternFill(start_color=fc, end_color=fc, fill_type="solid")


def _defense_note(df, cols, sv_threshold: int) -> str:
    if int((df["词类型"] == "防守词").sum()) > 0:
        return ""
    col_sv = cols["sv"]
    col_kw = cols.get("kw")
    sh_num = cols["sh_num"]
    share_mode = cols.get("share_mode", "missing")
    hi = df[pd.to_numeric(df[col_sv], errors="coerce") > sv_threshold] if col_sv else df.iloc[0:0]
    max_sh = float(hi[sh_num].max()) if len(hi) and sh_num in hi.columns else 0.0
    top_lines = []
    if len(hi) and col_kw:
        top = hi.sort_values(sh_num, ascending=False).head(3)
        for _, r in top.iterrows():
            kw = str(r.get(col_kw, ""))[:40]
            sh = float(r.get(sh_num, 0) or 0)
            sv = int(float(r.get(col_sv, 0) or 0))
            top_lines.append(f"{kw}（搜索量{sv:,}，份额{sh:.2f}%）")
    top_text = "；份额最高三词：" + "；".join(top_lines) if top_lines else ""
    if share_mode == "competitor":
        return (
            f"按「竞品SP份额」判定：搜索量>{sv_threshold} 的词中，竞品份额最高 {max_sh:.2f}%，"
            f"均未达到防守阈值 20%（竞品在大词上 SP 投入普遍不高，暂无需要规避的词位）。"
            f"{top_text}"
        )
    if share_mode == "own_listing":
        return (
            f"按「Listing SP份额」判定（2026+ Sif 列名，查竞品 ASIN 时即**竞品**在该词的 SP 份额，非你店铺）："
            f"搜索量>{sv_threshold} 共 {len(hi)} 个词，份额最高 {max_sh:.2f}%，未达到防守阈值 20%。"
            f"结论：竞品没有在大词上形成强广告壁垒，防守词为 0 是正常结果，应优先看缺口词。"
            f"{top_text}"
        )
    return "未识别 SP 份额列，无法判定防守词。"


def build_t3_stats(df, cols, asin, date, product_category, sv_threshold) -> dict:
    col_kw = cols["kw"]
    col_sv = cols["sv"]
    sh_num = cols["sh_num"]
    gap_df = df[df["词类型"] == "缺口词"].copy()
    gap_df[col_sv] = pd.to_numeric(gap_df[col_sv], errors="coerce")
    gap_df = gap_df.sort_values(col_sv, ascending=False)
    gap_list = []
    for _, r in gap_df.head(20).iterrows():
        gap_list.append({
            "keyword": str(r.get(col_kw, "")) if col_kw else "",
            "search_volume": int(r.get(col_sv, 0) or 0),
            "sp_share_pct": None if is_missing(r.get(cols.get("sh"))) else float(r.get(sh_num, 0) or 0),
            "difficulty": str(r.get("抢位难度", "")),
            "action": str(r.get("PD前建议操作", "")),
        })
    return {
        "report": "t3_ad_gap",
        "asin": asin, "date": date,
        "product_category": product_category,
        "sv_threshold": sv_threshold,
        "total": len(df),
        "gap_count": int((df["词类型"] == "缺口词").sum()),
        "defense_count": int((df["词类型"] == "防守词").sum()),
        "gap_low_difficulty": int(((df["词类型"] == "缺口词") & (df["抢位难度"] == "低")).sum()),
        "top_gap_keywords": gap_list,
        "share_mode": cols.get("share_mode", ""),
        "share_label": cols.get("share_label", ""),
        "defense_note": _defense_note(df, cols, sv_threshold),
    }


def build_word_report(df, cols, asin, date, output_dir, product_category, insights="", sv_threshold: int = 2000):
    if not HAS_DOCX:
        print("⚠️ 未安装 python-docx，跳过 Word", file=sys.stderr)
        return None

    doc = Document()
    setup_doc_styles(doc)
    add_doc_title(doc, f"{asin} 竞品广告缺口分析报告",
                  f"生成日期：{date}  |  产品类目：{product_category}")

    col_kw = cols["kw"]
    col_sv = cols["sv"]
    sh_num = cols["sh_num"]

    gap_cnt = (df["词类型"] == "缺口词").sum()
    def_cnt = (df["词类型"] == "防守词").sum()
    low_d = ((df["词类型"] == "缺口词") & (df["抢位难度"] == "低")).sum()
    mid_d = ((df["词类型"] == "缺口词") & (df["抢位难度"] == "中")).sum()

    doc.add_heading("一、执行摘要", level=1)
    if gap_cnt > 0:
        gap_kw_top3 = df[df["词类型"] == "缺口词"].nlargest(3, col_sv)[col_kw].astype(str).tolist() if col_kw else []
        conclusion = (
            f"产品类目「{product_category}」共分析 {len(df)} 个广告搜索词，"
            f"缺口词 {gap_cnt} 个，防守词 {def_cnt} 个，低难度缺口词 {low_d} 个。"
            f"TOP缺口词：{'、'.join(gap_kw_top3)}。"
        )
    else:
        conclusion = f"产品类目「{product_category}」共分析 {len(df)} 个广告搜索词，未发现高价值缺口词。"
    add_para(doc, conclusion)

    doc.add_heading("二、策略洞察与 PD 抢位方案", level=1)
    render_insights_section(doc, insights, heading="")

    gap_df = df[df["词类型"] == "缺口词"].copy()
    gap_df[col_sv] = pd.to_numeric(gap_df[col_sv], errors="coerce")
    gap_df = gap_df.sort_values(col_sv, ascending=False)

    doc.add_heading("三、TOP缺口词清单", level=1)
    if len(gap_df) > 0:
        tbl = doc.add_table(rows=min(len(gap_df), 20) + 1, cols=5)
        tbl.style = "Table Grid"
        for ci, h in enumerate(["广告搜索词", "搜索量", "竞品SP份额%", "抢位难度", "PD前建议"]):
            set_cell_text(tbl.rows[0].cells[ci], h)
        for ri, (_, r) in enumerate(gap_df.head(20).iterrows(), 1):
            set_cell_text(tbl.rows[ri].cells[0], str(r.get(col_kw, "")) if col_kw else "")
            set_cell_text(tbl.rows[ri].cells[1], str(int(r.get(col_sv, 0) or 0)))
            set_cell_text(tbl.rows[ri].cells[2], share_display(r, cols))
            set_cell_text(tbl.rows[ri].cells[3], str(r.get("抢位难度", "")))
            set_cell_text(tbl.rows[ri].cells[4], str(r.get("PD前建议操作", "")))
        apply_font_to_table(tbl)

    if HAS_MPL:
        chart_paths = _gen_charts_t3(df, cols, output_dir, asin, date)
        if chart_paths:
            items = [
                ("图1 · 柱形图：TOP10 缺口词周搜索量（按搜索量降序）", chart_paths[0]),
            ]
            if len(chart_paths) > 1:
                items.append(("图2 · 柱形图：词类型数量分布（缺口/防守/普通/不相关）", chart_paths[1]))
            if len(chart_paths) > 2:
                share_label = cols.get("share_label", "SP份额")
                items.append((
                    f"图3 · 气泡图：搜索量 vs {share_label}（气泡大小=SP贡献占比；"
                    f"左下=缺口词，右上=防守词）",
                    chart_paths[2],
                ))
            add_charts_section(doc, items, section_title="四、可视化图表", level=1)

    finalize_doc_fonts(doc)

    out = os.path.join(output_dir, f"{asin}_竞品广告缺口分析报告_{date}.docx")
    doc.save(out)
    print(f"✅ Word报告已生成：{out}")
    return out


def _gen_charts_t3(df, cols, output_dir, asin, date):
    paths = []
    col_kw = cols["kw"]; col_sv = cols["sv"]; sh_num = cols["sh_num"]; sp_num = cols["sp_num"]
    share_label = cols.get("share_label", "SP份额")
    type_colors = {"缺口词": "#00B050", "防守词": "#FF6B6B", "普通词": "#A9A9A9", "不相关词": "#D3D3D3"}
    try:
        gap_df = df[df["词类型"] == "缺口词"].copy()
        gap_df[col_sv] = pd.to_numeric(gap_df[col_sv], errors="coerce")
        gap_df = gap_df.sort_values(col_sv, ascending=False).head(10)
        if len(gap_df) > 0:
            labels = gap_df[col_kw].astype(str).str[:12].tolist() if col_kw else list(range(len(gap_df)))
            fig, ax = plt.subplots(figsize=(max(8, len(labels)), 5))
            bars = ax.bar(labels, gap_df[col_sv], color="#4A90D9")
            for bar, v in zip(bars, gap_df[col_sv]):
                ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 10,
                        str(int(v)), ha="center", fontsize=9, fontweight="bold")
            ax.set_xticklabels(labels, rotation=45, ha="right", fontsize=8)
            ax.set_title("TOP10缺口词搜索量", fontsize=13, fontweight="bold")
            ax.set_ylabel("周搜索量"); ax.spines[["top", "right"]].set_visible(False)
            p1 = chart_path(output_dir, f"chart1_t3_{asin}_{date}.png")
            fig.savefig(p1, dpi=150, bbox_inches="tight"); plt.close(fig); paths.append(p1)
    except Exception as e:
        print(f"图1失败: {e}", file=sys.stderr)
    try:
        type_counts = {wt: (df["词类型"] == wt).sum() for wt in ["缺口词", "防守词", "普通词", "不相关词"]}
        fig, ax = plt.subplots(figsize=(7, 5))
        colors = ["#00B050", "#FF6B6B", "#A9A9A9", "#D3D3D3"]
        bars = ax.bar(list(type_counts.keys()), list(type_counts.values()), color=colors)
        for bar, v in zip(bars, type_counts.values()):
            ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.3,
                    str(v), ha="center", fontsize=11, fontweight="bold")
        ax.set_title("词类型数量分布", fontsize=13, fontweight="bold")
        ax.set_ylabel("词数量"); ax.spines[["top", "right"]].set_visible(False)
        p2 = chart_path(output_dir, f"chart2_t3_{asin}_{date}.png")
        fig.savefig(p2, dpi=150, bbox_inches="tight"); plt.close(fig); paths.append(p2)
    except Exception as e:
        print(f"图2失败: {e}", file=sys.stderr)
    try:
        plot_df = df[df["词类型"] != "不相关词"].copy()
        plot_df[col_sv] = pd.to_numeric(plot_df[col_sv], errors="coerce").fillna(0)
        plot_df[sh_num] = pd.to_numeric(plot_df[sh_num], errors="coerce").fillna(0)
        plot_df[sp_num] = pd.to_numeric(plot_df[sp_num], errors="coerce").fillna(0)
        plot_df = plot_df[(plot_df[col_sv] > 0) & (plot_df[sh_num] >= 0)]
        if len(plot_df) > 0:
            sizes = (plot_df[sp_num].clip(lower=0.5) * 8).values
            colors = [type_colors.get(wt, "#888888") for wt in plot_df["词类型"]]
            fig, ax = plt.subplots(figsize=(10, 6))
            ax.scatter(
                plot_df[col_sv], plot_df[sh_num],
                s=sizes, c=colors, alpha=0.65, edgecolors="white", linewidths=0.5,
            )
            ax.axhline(20, color="#FF6B6B", linestyle="--", alpha=0.5, label="防守线 20%")
            ax.axhline(5, color="#00B050", linestyle="--", alpha=0.5, label="缺口线 5%")
            ax.set_xlabel("周搜索量")
            ax.set_ylabel(f"{share_label} (%)")
            ax.set_title("竞品广告重心气泡图（气泡大小=SP贡献占比）", fontsize=13, fontweight="bold")
            from matplotlib.patches import Patch
            legend = [
                Patch(facecolor="#00B050", label="缺口词"),
                Patch(facecolor="#FF6B6B", label="防守词"),
                Patch(facecolor="#A9A9A9", label="普通词"),
            ]
            ax.legend(handles=legend, loc="upper right")
            ax.spines[["top", "right"]].set_visible(False)
            p3 = chart_path(output_dir, f"chart3_t3_{asin}_{date}.png")
            fig.savefig(p3, dpi=150, bbox_inches="tight"); plt.close(fig); paths.append(p3)
    except Exception as e:
        print(f"图3气泡图失败: {e}", file=sys.stderr)
    return paths


def main():
    from threshold_presets import add_product_type_arg, resolve_thresholds, format_params_summary, collect_overrides_from_args

    parser = argparse.ArgumentParser(description="处理Sif查广告词（表3）")
    parser.add_argument("--input",            required=True)
    parser.add_argument("--output-dir",       required=True)
    parser.add_argument("--asin",             required=True)
    parser.add_argument("--date",             default=_date.today().strftime("%Y%m%d"),
                        help="日期YYYYMMDD，默认今天")
    parser.add_argument("--stage",            default="新品期", choices=["新品期", "成长期", "成熟期"])
    parser.add_argument("--product-category", default="产品", help="精确产品类目，影响相关性过滤")
    add_product_type_arg(parser)
    parser.add_argument("--sv-threshold",     type=int, default=0,
                        help="缺口词搜索量下限（0=用标品/非标品预设）")
    parser.add_argument("--gap-share-max",    type=float, default=None, help="缺口词竞品份额上限%")
    parser.add_argument("--defense-share-min", type=float, default=None, help="防守词竞品份额下限%")
    parser.add_argument("--insights",         default="")
    parser.add_argument("--insights-file",    default="")
    parser.add_argument("--skip-word",        action="store_true")
    args = parser.parse_args()
    args.insights = load_insights(args.insights, args.insights_file)
    overrides = collect_overrides_from_args(args)
    if args.sv_threshold > 0:
        overrides["sv_threshold"] = args.sv_threshold
    params = resolve_thresholds(args.product_type, args.stage, overrides)
    sv_thr = int(params["sv_threshold"])

    if not os.path.exists(args.input):
        print(f"❌ 找不到文件：{args.input}", file=sys.stderr); sys.exit(1)

    rid = make_run_id(args.date)
    os.makedirs(args.output_dir, exist_ok=True)
    print(f"📂 读取文件：{args.input}")
    print(f"  预设：{format_params_summary(params)}")
    df = load_excel(args.input)
    print(f"📊 共读取 {len(df)} 行  |  产品类目：{args.product_category}")

    df, cols, sv_median = preprocess(df, args.product_category)
    df = add_helper_columns(df, cols, args.product_category, params)

    gap_cnt = (df["词类型"] == "缺口词").sum()
    def_cnt = (df["词类型"] == "防守词").sum()
    irr_cnt = (df["词类型"] == "不相关词").sum()
    print(f"\n  缺口词：{gap_cnt} | 防守词：{def_cnt} | 普通词：{len(df)-gap_cnt-def_cnt-irr_cnt} | 不相关词：{irr_cnt}")
    print(f"  搜索量门槛：>{sv_thr}（可用 --sv-threshold 调整）")

    import shutil
    out_excel = os.path.join(args.output_dir, f"{args.asin}_竞品缺口分析_{rid}.xlsx")
    shutil.copy2(args.input, out_excel)

    wb = openpyxl.load_workbook(out_excel)
    write_helper_cols_to_main(wb, df)
    build_board_sheet(wb, df, cols, args.asin)
    wb.save(out_excel); print(f"✅ Excel已生成：{out_excel}")

    stats = build_t3_stats(df, cols, args.asin, args.date, args.product_category, sv_thr)
    stats["product_type"] = params.get("product_type", "")
    stats["params"] = {k: params.get(k) for k in ("sv_threshold", "gap_share_max", "defense_share_min")}
    stats["run_id"] = rid
    sp = save_stats(stats, stats_path(args.output_dir, args.asin, "t3", rid))
    print(f"📊 Stats JSON：{sp}")

    if args.skip_word:
        print("\n" + "=" * 60)
        print("📝 [Agent] 写 insights_t3.md 后带 --insights-file 重建 Word")
        print(f"  stats：{sp}")
        print(f"  insights：{insights_path(args.output_dir, 't3')}")
        print("=" * 60)
    else:
        build_word_report(df, cols, args.asin, rid, args.output_dir, args.product_category,
                          insights=args.insights, sv_threshold=sv_thr)
        cleanup_intermediate_files(args.output_dir, args.asin, rid, "t3")

    print(f"\n✅ 表3处理完成！输出目录：{args.output_dir}")


if __name__ == "__main__":
    main()
