"""
cross_analysis.py - 三表交叉分析，生成PD主攻词单
Usage:
  python cross_analysis.py \
    --t1 "关键词分层分析_20260611.xlsx" \
    --t2 "竞品弱点分析_20260611.xlsx" \
    --t3 "竞品缺口分析_20260611.xlsx" \
    --output-dir "处理结果/" \
    --asin B0CRMP3RQT \
    --date 20260611
"""
import argparse
import os
import sys
import warnings
warnings.filterwarnings("ignore")

from report_utils import (
    chart_path, get_col, pct_to_float, stats_path, save_stats, load_insights, cleanup_intermediate_files,
    setup_doc_styles, add_para, apply_font_to_table, set_cell_text, finalize_doc_fonts,
    render_insights_section, add_doc_title, add_picture_centered, add_charts_section, insights_path,
    is_missing, fmt_share_pct, fmt_cell,
)
from run_context import run_id as make_run_id

try:
    import pandas as pd
    import numpy as np
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.chart import BarChart, Reference, Series
except ImportError as e:
    print(f"缺少依赖：{e}\n请运行：pip install pandas openpyxl numpy", file=sys.stderr)
    sys.exit(1)

try:
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    plt.rcParams["font.family"] = ["Microsoft YaHei", "SimHei", "DejaVu Sans"]
    plt.rcParams["axes.unicode_minus"] = False
    HAS_MPL = True
except ImportError:
    HAS_MPL = False

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

THIN = Side(style="thin")
THIN_BORDER = Border(left=THIN, right=THIN, top=THIN, bottom=THIN)


def normalize_kw(kw: str) -> str:
    return str(kw).lower().strip()


# ──────── 从三张处理后的Excel读取关键词集合 ────────

def _read_sheet_with_header(path: str, sheet: str) -> pd.DataFrame:
    """自动检测表头行（兼容旧版 header=1 与新版 header=0）。
    覆盖三张表的特征列：T1(词级别/关键词)、T2(是否增长词/抢位机会评级)、T3(广告搜索词/词类型)
    """
    DETECT_COLS = [
        "词级别", "Level", "关键词", "Keyword",
        "抢位机会评级", "词库角色", "是否增长词",
        "广告搜索词", "词类型", "相关性判断",
    ]
    for header in (0, 1):
        df = pd.read_excel(path, sheet_name=sheet, header=header)
        if get_col(df, DETECT_COLS):
            return df
    return pd.read_excel(path, sheet_name=sheet, header=1)


def _read_main_sheet(path: str, skip_sheets: set) -> pd.DataFrame:
    """从处理后的 Excel 读取主数据 Sheet（跳过看板/汇总类 Sheet）。"""
    xl = pd.ExcelFile(path)
    for sheet in xl.sheet_names:
        if sheet in skip_sheets:
            continue
        df = _read_sheet_with_header(path, sheet)
        if len(df) == 0:
            continue
        return df
    return _read_sheet_with_header(path, xl.sheet_names[0])


def load_t1_sa(path: str):
    """读取表1处理结果，提取S级+A级词（从含辅助列的主数据 Sheet）。
    返回 (dataframe, col_kw, col_bid, col_conc, col_search)
    """
    xl = pd.ExcelFile(path)
    skip = {"分层汇总"}
    dfs = []
    for sheet in xl.sheet_names:
        if sheet in skip:
            continue
        df = _read_sheet_with_header(path, sheet)
        if get_col(df, ["词级别", "Level"]):
            dfs.append(df)
    if not dfs:
        df = _read_main_sheet(path, skip)
        dfs = [df]

    df = pd.concat(dfs, ignore_index=True)
    col_level  = get_col(df, ["词级别", "Level"])
    col_kw     = get_col(df, ["关键词", "Keyword"])
    col_bid    = get_col(df, ["建议竞价(中)", "建议竞价中值", "竞价中值", "Bid"])
    col_conc   = get_col(df, ["Top3点击集中度", "Top3转化集中度", "Top3集中度", "集中度", "Concentration"])
    col_search = get_col(df, ["周搜索量", "搜索量", "Search Volume"])
    if not col_level or not col_kw:
        raise ValueError(f"表1 {path} 找不到词级别或关键词列")
    sa = df[df[col_level].isin(["S级", "A级"])].copy()
    sa["__src"] = "T1"
    sa["__kw_norm"] = sa[col_kw].apply(normalize_kw)
    return sa, col_kw, col_bid, col_conc, col_search


def load_t2_high(path: str) -> pd.DataFrame:
    """读取表2处理结果，从主数据 Sheet 提取高机会词。"""
    df = _read_main_sheet(path, {"竞品弱点看板"})
    col_kw     = get_col(df, ["关键词", "Keyword"])
    col_rating = get_col(df, ["抢位机会评级", "评级", "Rating"])

    if col_rating and col_kw:
        high = df[df[col_rating].astype(str).str.strip() == "高"].copy()
    elif col_kw:
        high = df.copy()
    else:
        return pd.DataFrame(), None

    high["__src"] = "T2"
    high["__kw_norm"] = high[col_kw].apply(normalize_kw)
    return high, col_kw


def load_t3_gap(path: str) -> pd.DataFrame:
    """读取表3处理结果，从主数据 Sheet 提取缺口词。"""
    df = _read_main_sheet(path, {"竞品缺口看板"})
    col_kw   = get_col(df, ["广告搜索词", "关键词", "Keyword"])
    col_type = get_col(df, ["词类型", "Type"])

    if col_type and col_kw:
        gap = df[df[col_type].astype(str).str.strip() == "缺口词"].copy()
    elif col_kw:
        gap = df.copy()
    else:
        return pd.DataFrame(), None

    gap["__src"] = "T3"
    gap["__kw_norm"] = gap[col_kw].apply(normalize_kw)
    return gap, col_kw


# ──────── 交叉分析核心 ────────

def cross_analyze(t1_df, t1_kw, t2_df, t2_kw, t3_df, t3_kw,
                  t1_col_bid=None, t1_col_conc=None, t1_col_search=None):
    set1 = set(t1_df["__kw_norm"].dropna()) if len(t1_df) > 0 else set()
    set2 = set(t2_df["__kw_norm"].dropna()) if len(t2_df) > 0 else set()
    set3 = set(t3_df["__kw_norm"].dropna()) if len(t3_df) > 0 else set()

    sss_set = set1 & set2 & set3
    ss_12   = (set1 & set2) - sss_set
    ss_13   = (set1 & set3) - sss_set
    ss_23   = (set2 & set3) - sss_set
    ss_set  = ss_12 | ss_13 | ss_23
    s_set   = (set1 | set2 | set3) - sss_set - ss_set

    print(f"\n📊 三表交叉结果：")
    print(f"  表1(S+A级)：{len(set1)}词  表2(高机会)：{len(set2)}词  表3(缺口词)：{len(set3)}词")
    print(f"  SSS级（三表同时出现）：{len(sss_set)}词")
    print(f"  SS级（任意两表出现）：{len(ss_set)}词")
    print(f"  S级（仅单表出现）：{len(s_set)}词")

    if len(sss_set) == 0:
        print("  ⚠️  SSS级词为0，以SS级作为最终主攻词")

    # 反查每个词的详细数据
    t1_index = t1_df.set_index("__kw_norm") if len(t1_df) > 0 else pd.DataFrame()
    t2_index = t2_df.set_index("__kw_norm") if len(t2_df) > 0 else pd.DataFrame()
    t3_index = t3_df.set_index("__kw_norm") if len(t3_df) > 0 else pd.DataFrame()

    def get_word_row(kw, index_df, kw_col):
        if len(index_df) == 0 or kw_col is None:
            return {}
        try:
            rows = index_df.loc[[kw]] if kw in index_df.index else pd.DataFrame()
            if len(rows) == 0:
                return {}
            return rows.iloc[0].to_dict()
        except Exception:
            return {}

    def build_record(kw_norm, grade, sources):
        r1 = get_word_row(kw_norm, t1_index, t1_kw)
        r2 = get_word_row(kw_norm, t2_index, t2_kw)
        r3 = get_word_row(kw_norm, t3_index, t3_kw)

        # 原始关键词（取第一个有值的）
        orig_kw = (r1.get(t1_kw) or r2.get(t2_kw) or r3.get(t3_kw) or kw_norm) if (t1_kw or t2_kw or t3_kw) else kw_norm

        # 搜索量：优先用 t1 检测到的实际列名，再fallback
        sv = None
        sv_candidates = [t1_col_search] if t1_col_search else []
        sv_candidates += ["周搜索量", "搜索量", "Search Volume"]
        for c in sv_candidates:
            if c and c in r1 and pd.notna(r1[c]):
                try: sv = int(float(r1[c])); break
                except: pass
        if sv is None:
            for c in ["搜索量", "周搜索量"]:
                if c in r3 and pd.notna(r3[c]):
                    try: sv = int(float(r3[c])); break
                    except: pass

        # 竞价中值：优先用 t1 检测到的实际列名
        bid = None
        bid_candidates = [t1_col_bid] if t1_col_bid else []
        bid_candidates += ["建议竞价(中)", "建议竞价中值", "竞价中值", "Bid"]
        for c in bid_candidates:
            if c and c in r1 and pd.notna(r1[c]):
                try: bid = round(float(r1[c]), 2); break
                except: pass

        # 集中度：优先用 t1 检测到的实际列名
        conc = None
        conc_candidates = [t1_col_conc] if t1_col_conc else []
        conc_candidates += ["Top3点击集中度", "Top3转化集中度", "Top3集中度", "集中度", "Concentration"]
        for c in conc_candidates:
            if c and c in r1 and pd.notna(r1[c]):
                try: conc = round(float(r1[c]), 3); break
                except: pass

        # 竞品自然位状态
        nat_status = r2.get("竞品自然位状态", "") or ""

        # 竞品SP份额（__sh_num 不写Excel，改为从原始列名模糊匹配）
        sp_share = None
        for c in list(r3.keys()):
            if ("SP" in str(c) and "份额" in str(c)) or c == "__sh_num":
                try:
                    raw = r3[c]
                    if not is_missing(raw):
                        sp_share = round(pct_to_float(raw), 2)
                        break
                except Exception:
                    pass

        # 词级别（来自T1）
        word_level = r1.get("词级别", "") or ""

        # 匹配建议
        match_map = {"SSS级": "精准匹配（预算优先）", "SS级": "词组匹配+精准测试", "S级": "词组或广泛匹配"}
        match_suggest = match_map.get(grade, "")

        # PD加码优先级
        pd_priority = {"SSS级": "极高", "SS级": "高", "S级": "中"}.get(grade, "")

        # 综合操作建议
        op_map = {
            "SSS级": "PD前单独开精准匹配广告活动，预算优先倾斜",
            "SS级": "词组匹配主打，精准匹配测试",
            "S级": "词组或广泛匹配跟进，控制预算观察",
        }
        op = op_map.get(grade, "")

        return {
            "词级别": grade,
            "关键词": orig_kw,
            "出现来源": "+".join(sources),
            "周搜索量": sv,
            "建议竞价中值": bid,
            "Top3集中度": conc,
            "T1词级别": word_level,
            "竞品自然位状态": nat_status,
            "竞品SP份额%": sp_share,
            "匹配建议": match_suggest,
            "PD加码优先级": pd_priority,
            "综合操作建议": op,
        }

    records = []
    for kw in sss_set:
        records.append(build_record(kw, "SSS级", ["表1", "表2", "表3"]))
    for kw in ss_12:
        records.append(build_record(kw, "SS级", ["表1", "表2"]))
    for kw in ss_13:
        records.append(build_record(kw, "SS级", ["表1", "表3"]))
    for kw in ss_23:
        records.append(build_record(kw, "SS级", ["表2", "表3"]))
    for kw in s_set:
        src = []
        if kw in set1: src.append("表1")
        if kw in set2: src.append("表2")
        if kw in set3: src.append("表3")
        records.append(build_record(kw, "S级", src))

    result_df = pd.DataFrame(records)
    # 按级别排序（SSS > SS > S），再按搜索量排序
    grade_order = {"SSS级": 0, "SS级": 1, "S级": 2}
    result_df["__grade_order"] = result_df["词级别"].map(grade_order)
    result_df["__sv_sort"] = pd.to_numeric(result_df["周搜索量"], errors="coerce").fillna(0)
    result_df = result_df.sort_values(["__grade_order", "__sv_sort"], ascending=[True, False])
    result_df = result_df.drop(columns=["__grade_order", "__sv_sort"])

    overlap_meta = {
        "set1_count": len(set1),
        "set2_count": len(set2),
        "set3_count": len(set3),
        "ss_12_count": len(ss_12),
        "ss_13_count": len(ss_13),
        "ss_23_count": len(ss_23),
        "union_count": len(set1 | set2 | set3),
        "sss_keywords": [build_record(k, "SSS级", ["表1", "表2", "表3"])["关键词"] for k in list(sss_set)[:5]],
        "ss_12_samples": [build_record(k, "SS级", ["表1", "表2"])["关键词"] for k in list(ss_12)[:5]],
        "ss_13_samples": [build_record(k, "SS级", ["表1", "表3"])["关键词"] for k in list(ss_13)[:5]],
        "ss_23_samples": [build_record(k, "SS级", ["表2", "表3"])["关键词"] for k in list(ss_23)[:5]],
        "s_only_t1_count": len(set1 - set2 - set3),
        "s_only_t2_count": len(set2 - set1 - set3),
        "s_only_t3_count": len(set3 - set1 - set2),
        "naming_note": "交叉S级=仅出现在一张筛选表，与表1 S级不同；主攻预算应聚焦SSS+SS",
    }

    return result_df, len(sss_set), len(ss_set), len(s_set), overlap_meta


# ──────── 生成PD主攻词单Excel ────────

def build_output_excel(result_df, asin, date, output_dir):
    out_path = os.path.join(output_dir, f"{asin}_PD主攻词单_{date}.xlsx")
    wb = openpyxl.Workbook()

    title_font = Font(bold=True, name="Microsoft YaHei", size=13)
    head_fill  = PatternFill(start_color="1F4E79", end_color="1F4E79", fill_type="solid")
    head_font  = Font(color="FFFFFF", bold=True, name="Microsoft YaHei", size=10)

    grade_fills = {
        "SSS级": PatternFill(start_color="C6EFCE", end_color="C6EFCE", fill_type="solid"),
        "SS级":  PatternFill(start_color="FFEB9C", end_color="FFEB9C", fill_type="solid"),
        "S级":   PatternFill(start_color="FCE4D6", end_color="FCE4D6", fill_type="solid"),
    }

    output_cols = ["词级别", "关键词", "出现来源", "周搜索量", "建议竞价中值",
                   "Top3集中度", "竞品自然位状态", "竞品SP份额%", "匹配建议", "PD加码优先级", "综合操作建议"]
    col_widths  = [10, 25, 14, 12, 14, 12, 16, 14, 20, 14, 36]

    # Sheet 1: 按级别分开
    for grade in ["SSS级", "SS级", "S级"]:
        sub = result_df[result_df["词级别"] == grade]
        ws = wb.create_sheet(f"{grade}词（{'三表交叉' if grade=='SSS级' else '两表交叉' if grade=='SS级' else '单表出现'}）")
        ws.sheet_view.showGridLines = False
        # 表头
        for ci, (col, w) in enumerate(zip(output_cols, col_widths), 1):
            c = ws.cell(1, ci, col)
            c.fill = head_fill; c.font = head_font; c.border = THIN_BORDER
            c.alignment = Alignment(horizontal="center")
            ws.column_dimensions[openpyxl.utils.get_column_letter(ci)].width = w
        # 数据
        for ri, (_, row) in enumerate(sub.iterrows(), 2):
            fill = grade_fills.get(grade) if ri % 2 == 0 else None
            for ci, col in enumerate(output_cols, 1):
                val = row.get(col, "")
                if col == "竞品SP份额%":
                    display = fmt_share_pct(val)
                elif pd.notna(val):
                    display = val
                else:
                    display = ""
                c = ws.cell(ri, ci, display)
                c.border = THIN_BORDER
                c.font = Font(name="Microsoft YaHei", size=9)
                c.alignment = Alignment(wrap_text=(col == "综合操作建议"))
                if fill: c.fill = fill

    # Sheet 4: 词级统计（带图表）
    ws_stats = wb.create_sheet("词级统计")
    ws_stats.sheet_view.showGridLines = False
    ws_stats.cell(1, 1, "📊 PD主攻词单统计汇总").font = title_font

    stats_data = []
    total_all = len(result_df)
    for grade in ["SSS级", "SS级", "S级"]:
        cnt = (result_df["词级别"] == grade).sum()
        sub = result_df[result_df["词级别"] == grade]
        avg_sv  = pd.to_numeric(sub["周搜索量"], errors="coerce").mean()
        stats_data.append((grade, cnt, f"{cnt/total_all*100:.1f}%" if total_all else "0%",
                           f"{avg_sv:.0f}" if pd.notna(avg_sv) else "N/A"))

    stat_headers = ["词级别", "词数量", "占比", "平均搜索量"]
    for ci, h in enumerate(stat_headers, 1):
        c = ws_stats.cell(3, ci, h); c.fill = head_fill; c.font = head_font; c.border = THIN_BORDER
    for ri, (grade, cnt, pct, avg) in enumerate(stats_data, 4):
        ws_stats.cell(ri, 1, grade).border = THIN_BORDER
        ws_stats.cell(ri, 2, cnt).border   = THIN_BORDER
        ws_stats.cell(ri, 3, pct).border   = THIN_BORDER
        ws_stats.cell(ri, 4, avg).border   = THIN_BORDER
        gfill = grade_fills.get(grade)
        if gfill:
            for ci in range(1, 5):
                ws_stats.cell(ri, ci).fill = gfill

    # 嵌入柱状图
    chart_data_row = 10
    for i, (grade, cnt, _, _) in enumerate(stats_data):
        ws_stats.cell(chart_data_row + i, 8, grade)
        ws_stats.cell(chart_data_row + i, 9, cnt)
    bc = BarChart()
    bc.type = "col"; bc.title = "SSS/SS/S级词数量分布"; bc.y_axis.title = "词数量"
    bc.width = 14; bc.height = 9
    dr = Reference(ws_stats, min_col=9, min_row=chart_data_row, max_row=chart_data_row + 2)
    cr = Reference(ws_stats, min_col=8, min_row=chart_data_row, max_row=chart_data_row + 2)
    s1 = Series(dr, title="词数量"); bc.series.append(s1); bc.set_categories(cr)
    ws_stats.add_chart(bc, "A7")

    # 图二：SSS+SS级词搜索量 TOP10
    top_ss = result_df[result_df["词级别"].isin(["SSS级", "SS级"])].copy()
    top_ss["__sv"] = pd.to_numeric(top_ss["周搜索量"], errors="coerce").fillna(0)
    top_ss = top_ss.sort_values("__sv", ascending=False).head(10)
    sv_row = 14
    for i, (_, r) in enumerate(top_ss.iterrows()):
        kw = str(r.get("关键词", ""))[:8]
        ws_stats.cell(sv_row + i, 8, kw)
        ws_stats.cell(sv_row + i, 9, int(r.get("__sv", 0)))
    if len(top_ss) > 0:
        bc2 = BarChart()
        bc2.type = "col"
        bc2.title = "SSS+SS级词搜索量 TOP10"
        bc2.y_axis.title = "周搜索量"
        bc2.width = 14
        bc2.height = 9
        dr2 = Reference(ws_stats, min_col=9, min_row=sv_row, max_row=sv_row + len(top_ss) - 1)
        cr2 = Reference(ws_stats, min_col=8, min_row=sv_row, max_row=sv_row + len(top_ss) - 1)
        bc2.series.append(Series(dr2, title="搜索量"))
        bc2.set_categories(cr2)
        ws_stats.add_chart(bc2, "J7")

    # 删除默认空Sheet
    if "Sheet" in wb.sheetnames:
        del wb["Sheet"]

    wb.save(out_path)
    print(f"✅ PD主攻词单Excel已生成：{out_path}")
    return out_path


# ──────── 生成Word报告 ────────

def build_cross_stats(result_df, sss_cnt, ss_cnt, s_cnt, asin, date, overlap_meta=None) -> dict:
    top = result_df[result_df["词级别"].isin(["SSS级", "SS级"])].head(15)
    records = []
    for _, r in top.iterrows():
        records.append({
            "grade": str(r.get("词级别", "")),
            "keyword": str(r.get("关键词", "")),
            "sources": str(r.get("出现来源", "")),
            "search_volume": r.get("周搜索量"),
            "bid": r.get("建议竞价中值"),
            "concentration": r.get("Top3集中度"),
            "sp_share": r.get("竞品SP份额%"),
            "action": str(r.get("综合操作建议", "")),
        })
    stats = {
        "report": "cross_pd_primary",
        "asin": asin, "date": date,
        "sss_count": sss_cnt, "ss_count": ss_cnt, "s_count": s_cnt,
        "primary_focus_count": sss_cnt + ss_cnt,
        "primary_keywords": records,
    }
    if overlap_meta:
        stats["overlap"] = overlap_meta
    return stats


def build_word_report(result_df, sss_cnt, ss_cnt, s_cnt, asin, date, output_dir, insights="", overlap_meta=None):
    if not HAS_DOCX:
        print("⚠️ 未安装 python-docx，跳过 Word", file=sys.stderr)
        return None

    doc = Document()
    setup_doc_styles(doc)
    add_doc_title(doc, f"{asin} PD主攻词单分析报告", f"生成日期：{date}")

    doc.add_heading("一、执行摘要", level=1)
    main_grade = "SSS级" if sss_cnt > 0 else "SS级"
    main_cnt = sss_cnt if sss_cnt > 0 else ss_cnt
    top_words = result_df[result_df["词级别"].isin(["SSS级", "SS级"])].head(5)
    top_kw_str = "、".join(top_words["关键词"].astype(str).tolist()) if len(top_words) else "（无）"
    top_sv = top_words["周搜索量"].apply(pd.to_numeric, errors="coerce")
    avg_sv_str = f"{top_sv.mean():.0f}" if len(top_sv) > 0 and top_sv.notna().any() else "N/A"
    add_para(doc,
        f"三表交叉：SSS级 {sss_cnt} 个，SS级 {ss_cnt} 个，S级 {s_cnt} 个（单表词，非主攻核心）。"
        f"PD 主攻以{main_grade}为核心（{main_cnt}个），均值搜索量 {avg_sv_str}。TOP5：{top_kw_str}。"
    )

    if overlap_meta:
        doc.add_heading("二、交叉分级逻辑（数据校验）", level=1)
        add_para(doc,
            "说明：交叉「S级」= 只出现在一张筛选表中的词，与表1「S级」含义不同。"
            "S 数量偏多属正常——三表并集减去双表/三表交集后的余量。"
        )
        logic_tbl = doc.add_table(rows=8, cols=2)
        logic_tbl.style = "Table Grid"
        rows_data = [
            ("表1 输入（S+A级）", str(overlap_meta.get("set1_count", ""))),
            ("表2 输入（高机会）", str(overlap_meta.get("set2_count", ""))),
            ("表3 输入（缺口词）", str(overlap_meta.get("set3_count", ""))),
            ("三表并集（去重）", str(overlap_meta.get("union_count", ""))),
            ("SSS（三表交集）", str(sss_cnt)),
            ("SS 表1+表2 / 表1+表3 / 表2+表3",
             f"{overlap_meta.get('ss_12_count', 0)} / {overlap_meta.get('ss_13_count', 0)} / {overlap_meta.get('ss_23_count', 0)}"),
            ("S 仅单表（表1/表2/表3）",
             f"{overlap_meta.get('s_only_t1_count', 0)} / {overlap_meta.get('s_only_t2_count', 0)} / {overlap_meta.get('s_only_t3_count', 0)}"),
            ("校验公式", f"SSS+SS+S = {sss_cnt}+{ss_cnt}+{s_cnt} = {sss_cnt + ss_cnt + s_cnt}（应等于并集）"),
        ]
        for ri, (k, v) in enumerate(rows_data):
            set_cell_text(logic_tbl.rows[ri].cells[0], k)
            set_cell_text(logic_tbl.rows[ri].cells[1], v)
        apply_font_to_table(logic_tbl)

    doc.add_heading("三、PD 备战策略（AI 分析）", level=1)
    render_insights_section(doc, insights, heading="")

    sss_df = result_df[result_df["词级别"] == "SSS级"].head(10)
    sec = 4
    if sss_cnt > 0:
        doc.add_heading("四、SSS级词逐一解读", level=1)
        sec = 5
        for _, r in sss_df.iterrows():
            kw = r.get("关键词", "")
            sv = r.get("周搜索量", "N/A")
            bid = r.get("建议竞价中值", "N/A")
            conc = r.get("Top3集中度", "N/A")
            bid_str = f"${bid:.2f}" if isinstance(bid, float) else str(bid)
            conc_str = f"{conc:.3f}" if isinstance(conc, float) else str(conc)
            sp_str = fmt_share_pct(r.get("竞品SP份额%"))
            sp_part = f"，SP份额 {sp_str}" if sp_str else ""
            add_para(doc,
                f"【{kw}】搜索量 {sv}，竞价 {bid_str}，集中度 {conc_str}，"
                f"自然位 {r.get('竞品自然位状态', '未知')}{sp_part} "
                f"→ {r.get('综合操作建议', '')}"
            )

    ss_df = result_df[result_df["词级别"] == "SS级"].head(20)
    cn = ["四", "五", "六", "七"]
    sec_idx = sec - 4
    if len(ss_df) > 0:
        doc.add_heading(f"{cn[sec_idx]}、SS级词清单（主攻候选）", level=1)
        sec_idx += 1
        tbl = doc.add_table(rows=len(ss_df) + 1, cols=6)
        tbl.style = "Table Grid"
        for ci, h in enumerate(["关键词", "来源", "周搜索量", "竞品自然位", "竞品SP份额%", "操作建议"]):
            set_cell_text(tbl.rows[0].cells[ci], h)
        for ri, (_, r) in enumerate(ss_df.iterrows(), 1):
            set_cell_text(tbl.rows[ri].cells[0], str(r.get("关键词", "")))
            set_cell_text(tbl.rows[ri].cells[1], str(r.get("出现来源", "")))
            set_cell_text(tbl.rows[ri].cells[2], str(r.get("周搜索量", "")))
            set_cell_text(tbl.rows[ri].cells[3], str(r.get("竞品自然位状态", "")))
            set_cell_text(tbl.rows[ri].cells[4], fmt_share_pct(r.get("竞品SP份额%")))
            set_cell_text(tbl.rows[ri].cells[5], str(r.get("综合操作建议", "")))
        apply_font_to_table(tbl)

    if s_cnt > 0:
        doc.add_heading(f"{cn[sec_idx]}、S级单表词说明", level=1)
        add_para(doc,
            f"共 {s_cnt} 个词仅出现在一张筛选表中，不建议作为 PD 主攻预算核心。"
            f"完整清单见 Excel「PD主攻词单」Sheet，可按搜索量择优观察。"
        )

    if HAS_MPL:
        chart_paths = _gen_charts_cross(result_df, sss_cnt, ss_cnt, s_cnt, output_dir, asin, date)
        if chart_paths:
            items = [
                ("图1 · 柱形图：三表交叉分级词数量（SSS / SS / S）", chart_paths[0]),
            ]
            if len(chart_paths) > 1:
                items.append((
                    "图2 · 雷达图：SSS+SS 级词综合价值（搜索量/竞价友好/竞品防守弱/自然位空间/集中度）",
                    chart_paths[1],
                ))
            if len(chart_paths) > 2:
                items.append(("图3 · 柱形图：SS/SSS 级 TOP 词周搜索量", chart_paths[2]))
            add_charts_section(
                doc, items,
                section_title=f"{cn[min(sec_idx + 1, 3)]}、可视化图表",
                level=1,
            )

    finalize_doc_fonts(doc)
    out = os.path.join(output_dir, f"{asin}_PD主攻词单报告_{date}.docx")
    try:
        doc.save(out)
    except PermissionError:
        alt = os.path.join(output_dir, f"{asin}_PD主攻词单报告_{date}_new.docx")
        doc.save(alt)
        out = alt
        print(f"⚠️ 原 Word 被占用，已另存：{alt}", file=sys.stderr)
    print(f"✅ Word报告已生成：{out}")
    return out


def _gen_charts_cross(result_df, sss_cnt, ss_cnt, s_cnt, output_dir, asin, date):
    paths = []
    try:
        counts = [sss_cnt, ss_cnt, s_cnt]
        labels = ["SSS级", "SS级", "S级"]
        fig, ax = plt.subplots(figsize=(7, 5))
        colors = ["#70AD47", "#FFC000", "#FF6B6B"]
        bars = ax.bar(labels, counts, color=colors, width=0.4)
        for bar, cnt in zip(bars, counts):
            ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.2,
                    str(cnt), ha="center", fontsize=14, fontweight="bold")
        ax.set_title("SSS/SS/S级词数量分布", fontsize=13, fontweight="bold")
        ax.set_ylabel("词数量"); ax.spines[["top", "right"]].set_visible(False)
        p1 = chart_path(output_dir, f"chart1_cross_{asin}_{date}.png")
        fig.savefig(p1, dpi=150, bbox_inches="tight"); plt.close(fig); paths.append(p1)
    except Exception as e:
        print(f"图1失败: {e}", file=sys.stderr)

    try:
        core = result_df[result_df["词级别"].isin(["SSS级", "SS级"])].copy()
        if len(core) > 0:
            sv = pd.to_numeric(core["周搜索量"], errors="coerce").fillna(0)
            bid = pd.to_numeric(core["建议竞价中值"], errors="coerce").fillna(1.0)
            conc = pd.to_numeric(core["Top3集中度"], errors="coerce").fillna(0.35)
            sp = pd.to_numeric(core["竞品SP份额%"], errors="coerce").fillna(10.0)

            def norm_high(v, ref):
                return float(np.clip(v / ref * 100, 0, 100)) if ref else 50.0

            def norm_low(v, ref):
                return float(np.clip((1 - v / ref) * 100, 0, 100)) if ref else 50.0

            sv_score = norm_high(sv.median(), max(sv.max(), 1))
            bid_score = norm_low(bid.median(), max(bid.max(), 0.5))
            defense_score = norm_low(sp.median(), 25.0)
            conc_score = norm_low(conc.median(), max(conc.max(), 0.35))

            nat_map = {"纯广告": 90, "未稳": 65, "稳定": 30}
            nat_vals = core["竞品自然位状态"].astype(str).map(lambda x: nat_map.get(x.strip(), 50))
            nat_score = float(nat_vals.mean()) if len(nat_vals) else 50.0

            dims = ["搜索量", "竞价友好", "竞品防守弱", "自然位空间", "集中度友好"]
            values = [sv_score, bid_score, defense_score, nat_score, conc_score]
            angles = np.linspace(0, 2 * np.pi, len(dims), endpoint=False).tolist()
            values_cycle = values + values[:1]
            angles_cycle = angles + angles[:1]

            fig, ax = plt.subplots(figsize=(7, 7), subplot_kw=dict(polar=True))
            ax.plot(angles_cycle, values_cycle, "o-", linewidth=2, color="#2E75B6")
            ax.fill(angles_cycle, values_cycle, alpha=0.25, color="#2E75B6")
            ax.set_xticks(angles)
            ax.set_xticklabels(dims, fontsize=10)
            ax.set_ylim(0, 100)
            ax.set_title("SSS+SS级词综合价值雷达图", fontsize=13, fontweight="bold", pad=20)
            p2 = chart_path(output_dir, f"chart2_cross_{asin}_{date}.png")
            fig.savefig(p2, dpi=150, bbox_inches="tight"); plt.close(fig); paths.append(p2)
    except Exception as e:
        print(f"图2雷达图失败: {e}", file=sys.stderr)

    try:
        top_df = result_df[result_df["词级别"].isin(["SSS级", "SS级"])].head(10)
        top_df["__sv"] = pd.to_numeric(top_df["周搜索量"], errors="coerce").fillna(0)
        top_df = top_df.sort_values("__sv", ascending=False)
        if len(top_df) > 0:
            labels_t = top_df["关键词"].astype(str).str[:12].tolist()
            colors_t = ["#70AD47" if g == "SSS级" else "#FFC000" for g in top_df["词级别"]]
            fig, ax = plt.subplots(figsize=(max(9, len(labels_t)), 5))
            bars = ax.bar(labels_t, top_df["__sv"], color=colors_t)
            for bar, v in zip(bars, top_df["__sv"]):
                ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 20,
                        str(int(v)), ha="center", fontsize=8, fontweight="bold")
            ax.set_xticklabels(labels_t, rotation=45, ha="right", fontsize=8)
            ax.set_title("SSS+SS级词搜索量 TOP10", fontsize=12, fontweight="bold")
            ax.set_ylabel("周搜索量"); ax.spines[["top", "right"]].set_visible(False)
            from matplotlib.patches import Patch
            legend_elements = [Patch(facecolor="#70AD47", label="SSS级"), Patch(facecolor="#FFC000", label="SS级")]
            ax.legend(handles=legend_elements)
            p3 = chart_path(output_dir, f"chart3_cross_{asin}_{date}.png")
            fig.savefig(p3, dpi=150, bbox_inches="tight"); plt.close(fig); paths.append(p3)
    except Exception as e:
        print(f"图3失败: {e}", file=sys.stderr)
    return paths


def main():
    from datetime import date as _date, datetime
    parser = argparse.ArgumentParser(description="三表交叉分析，生成PD主攻词单")
    parser.add_argument("--t1",         required=True, help="表1处理结果Excel路径")
    parser.add_argument("--t2",         required=True, help="表2处理结果Excel路径")
    parser.add_argument("--t3",         required=True, help="表3处理结果Excel路径")
    parser.add_argument("--output-dir", required=True)
    parser.add_argument("--asin",       required=True)
    parser.add_argument("--date",       default=_date.today().strftime("%Y%m%d"))
    parser.add_argument("--insights",      default="")
    parser.add_argument("--insights-file", default="")
    parser.add_argument("--skip-word",     action="store_true")
    args = parser.parse_args()
    args.insights = load_insights(args.insights, args.insights_file)

    for p in [args.t1, args.t2, args.t3]:
        if not os.path.exists(p):
            print(f"❌ 找不到文件：{p}", file=sys.stderr); sys.exit(1)

    rid = make_run_id(args.date)
    os.makedirs(args.output_dir, exist_ok=True)

    print(f"📂 读取表1（关键词分层）：{args.t1}")
    t1_df, t1_kw, t1_bid, t1_conc, t1_sv = load_t1_sa(args.t1)
    print(f"   S+A级词：{len(t1_df)} 个  |  bid列：{t1_bid}  |  conc列：{t1_conc}")

    print(f"📂 读取表2（高机会词）：{args.t2}")
    t2_df, t2_kw = load_t2_high(args.t2)
    print(f"   高机会词：{len(t2_df)} 个")

    print(f"📂 读取表3（缺口词）：{args.t3}")
    t3_df, t3_kw = load_t3_gap(args.t3)
    print(f"   缺口词：{len(t3_df)} 个")

    result_df, sss_cnt, ss_cnt, s_cnt, overlap_meta = cross_analyze(
        t1_df, t1_kw, t2_df, t2_kw, t3_df, t3_kw,
        t1_col_bid=t1_bid, t1_col_conc=t1_conc, t1_col_search=t1_sv
    )

    print("\n📝 生成PD主攻词单Excel...")
    build_output_excel(result_df, args.asin, rid, args.output_dir)

    stats = build_cross_stats(result_df, sss_cnt, ss_cnt, s_cnt, args.asin, args.date, overlap_meta)
    stats["run_id"] = rid
    sp = save_stats(stats, stats_path(args.output_dir, args.asin, "cross", rid))
    print(f"📊 Stats JSON：{sp}")

    if args.skip_word:
        print("\n" + "=" * 60)
        print("📝 [Agent] 写 insights_cross.md 后带 --insights-file 重建 Word")
        print(f"  stats：{sp}")
        print(f"  insights：{insights_path(args.output_dir, 'cross')}")
        print("=" * 60)
    else:
        build_word_report(result_df, sss_cnt, ss_cnt, s_cnt,
                          args.asin, rid, args.output_dir, insights=args.insights,
                          overlap_meta=overlap_meta)
        cleanup_intermediate_files(args.output_dir, args.asin, rid, "cross")

    print(f"\n✅ 三表交叉分析完成！SSS:{sss_cnt} | SS:{ss_cnt} | S:{s_cnt}")


if __name__ == "__main__":
    main()
