"""
process_table1.py - 处理表1：关键词调研
Usage:
  # 仅计算基准（用户确认前）
  python process_table1.py --mode compute-thresholds --input "关键词调研_20260611.xlsx"

  # 完整处理
  python process_table1.py --mode process --input "..." --output-dir "..." --asin B0CRMP3RQT --date 20260611 --stage 新品期

  # 覆盖门槛
  python process_table1.py --mode process ... --s-threshold 5000 --a-threshold 1000
"""
import argparse
import os
import sys
import json
import warnings
warnings.filterwarnings("ignore")

from report_utils import (
    charts_dir, chart_path, get_col, stats_path, save_stats, load_insights, cleanup_intermediate_files,
    setup_doc_styles, add_para, apply_font_to_table, set_cell_text, finalize_doc_fonts,
    render_insights_section, add_doc_title, add_picture_centered, add_charts_section, llm_prompt_template, insights_path,
)
from run_context import run_id as make_run_id

try:
    import pandas as pd
    import numpy as np
    import openpyxl
    from openpyxl.styles import (Font, PatternFill, Alignment, Border, Side,
                                  numbers as openpyxl_numbers)
    from openpyxl.chart import BarChart, ScatterChart, Reference, Series
    from openpyxl.chart.label import DataLabelList
except ImportError as e:
    print(f"缺少依赖：{e}\n请运行：pip install pandas openpyxl numpy", file=sys.stderr)
    sys.exit(1)

try:
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    plt.rcParams["font.family"] = ["Microsoft YaHei", "SimHei", "DejaVu Sans"]
    plt.rcParams["axes.unicode_minus"] = False
    HAS_MPL = True
except ImportError:
    HAS_MPL = False

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

RELEVANT_SHEETS = ["高相关", "中相关", "低相关"]
THIN = Side(style="thin")
THIN_BORDER = Border(left=THIN, right=THIN, top=THIN, bottom=THIN)


# ───────────── 读取数据 ─────────────

def load_excel(path: str) -> pd.DataFrame:
    """读取相关 Sheet 并合并。兼容旧版（三个独立 Sheet + 第2行表头）和新版（合并 Sheet + 第1行表头）。"""
    dfs = []
    xl = pd.ExcelFile(path)
    available = xl.sheet_names

    # 旧版：高/中/低 三个独立 Sheet，第1行空、第2行表头
    for sheet in RELEVANT_SHEETS:
        if sheet in available:
            df = pd.read_excel(path, sheet_name=sheet, header=1)
            df["__sheet__"] = sheet
            df["__header_row__"] = 2
            dfs.append(df)

    # 新版（2026-06 Sif 8步向导）：合并 Sheet，第1行即表头
    if not dfs:
        combined = [s for s in available if "高" in s and "低" in s and "相关" in s]
        for sheet in combined:
            df = pd.read_excel(path, sheet_name=sheet, header=0)
            df["__sheet__"] = sheet
            df["__header_row__"] = 1
            dfs.append(df)

    if not dfs:
        raise ValueError(f"未找到任何相关Sheet（期望：{RELEVANT_SHEETS} 或含「高/低相关」合并Sheet），实际：{available}")
    return pd.concat(dfs, ignore_index=True)


# ───────────── 计算基准 ─────────────

def compute_thresholds(df: pd.DataFrame, s_percentile: int = 97, a_percentile: int = 85) -> dict:
    col_search = get_col(df, ["周搜索量", "搜索量", "Search Volume"])
    col_bid    = get_col(df, ["建议竞价(中)", "建议竞价中值", "竞价中值", "Bid"])
    col_conc   = get_col(df, ["Top3点击集中度", "Top3转化集中度", "Top3集中度", "集中度", "Concentration"])
    col_rel    = get_col(df, ["相关性", "Relevance"])

    if not col_search:
        raise ValueError("未找到周搜索量列，请检查表格结构")

    high_mask = df[col_rel].astype(str).str.contains("高") if col_rel else pd.Series([True] * len(df))
    high_df = df[high_mask].copy()

    search_vals = pd.to_numeric(high_df[col_search], errors="coerce").dropna()
    bid_vals    = pd.to_numeric(high_df[col_bid], errors="coerce").dropna() if col_bid else pd.Series(dtype=float)
    conc_vals   = pd.to_numeric(high_df[col_conc], errors="coerce").dropna() if col_conc else pd.Series(dtype=float)

    s_thresh = int(np.percentile(search_vals, s_percentile)) if len(search_vals) > 0 else 0
    a_thresh = int(np.percentile(search_vals, a_percentile)) if len(search_vals) > 0 else 0
    bid_med  = round(float(bid_vals.median()), 2) if len(bid_vals) > 0 else 0.0
    conc_med = round(float(conc_vals.median()), 3) if len(conc_vals) > 0 else 0.0

    return {
        "total_high": int(high_mask.sum()),
        "total_words": len(df),
        "s_percentile": s_percentile,
        "a_percentile": a_percentile,
        "s_threshold": s_thresh,
        "a_threshold": a_thresh,
        "bid_median": bid_med,
        "conc_median": conc_med,
        "col_search": col_search,
        "col_bid": col_bid,
        "col_conc": col_conc,
        "col_rel": col_rel,
    }


# ───────────── 分级逻辑 ─────────────

def classify_word(row, t: dict, s_threshold: int, a_threshold: int, stage: str,
                  params: dict | None = None) -> tuple:
    p = params or {}
    s_conc = float(p.get("s_conc_max", 0.30))
    a_conc = float(p.get("a_conc_max", 0.35))
    s_bid_mult = float(p.get("s_bid_mult", 1.2))
    a_bid_mult = float(p.get("a_bid_mult", 1.5))
    b_require_cvr = bool(p.get("b_require_cvr", True))
    rel   = str(row.get(t["col_rel"], "")).strip()
    try:
        search = float(row.get(t["col_search"], 0) or 0)
    except (ValueError, TypeError):
        search = 0.0
    try:
        bid = float(row.get(t["col_bid"], 0) or 0)
    except (ValueError, TypeError):
        bid = 0.0
    try:
        conc = float(row.get(t["col_conc"], 1) or 1)
    except (ValueError, TypeError):
        conc = 1.0
    try:
        cvr_raw = row.get(t.get("col_cvr", ""), None)
        cvr_valid = cvr_raw is not None and str(cvr_raw).strip() not in ("", "nan", "None")
    except Exception:
        cvr_valid = False

    bid_med = t["bid_median"]
    is_high = "高" in rel
    is_mid  = "中" in rel
    is_low  = "低" in rel

    # S 级
    if (is_high and search >= s_threshold
            and (bid_med == 0 or bid < bid_med * s_bid_mult)
            and conc < s_conc):
        pd_priority = "中（建议成熟期再主攻）" if stage == "新品期" else "高"
        return "S级", f"高相关+搜索量{int(search)}≥S级门槛{s_threshold}+竞价友好+集中度{conc:.2f}<{s_conc:.2f}", "精准匹配", pd_priority

    # A 级
    if is_high and a_threshold <= search < s_threshold and (bid_med == 0 or bid < bid_med * a_bid_mult) and conc < a_conc:
        return "A级", f"高相关+搜索量{int(search)}在A-S门槛间+集中度{conc:.2f}<{a_conc:.2f}", "词组匹配", "高"
    if is_high and search >= s_threshold and conc >= s_conc:
        return "A级", f"高相关+搜索量达S级但集中度{conc:.2f}≥{s_conc:.2f}降级", "词组匹配", "高"

    # B 级（兼容小类目 a_threshold < 200 的情况，用较小值作下限）
    b_low = min(200, a_threshold)
    if (is_high or is_mid) and b_low <= search < a_threshold and (cvr_valid or not b_require_cvr):
        cvr_note = "有转化数据" if cvr_valid else "新品期放宽转化率条件"
        return "B级", f"中高相关+搜索量{int(search)}在{b_low}-A门槛间+{cvr_note}", "广泛匹配", "中"
    # 中相关+大词（>=A门槛）但不满足高相关S/A条件 → B级观察（不排除，但匹配方式宽泛）
    if is_mid and search >= a_threshold:
        return "B级", f"中相关+搜索量{int(search)}≥A门槛，不纳入S/A，宽泛跟投观察转化", "广泛匹配", "低"

    # C 级
    reason_parts = []
    if is_low:
        reason_parts.append("低相关")
    if search < 200 and not cvr_valid:
        reason_parts.append(f"搜索量{int(search)}<200且无转化")
    if conc > 0.5:
        reason_parts.append(f"集中度{conc:.2f}>0.5")
    if not reason_parts:
        reason_parts.append("不满足S/A/B级条件")
    return "C级", "+".join(reason_parts), "否词", "不建议"


def get_competition(conc: float) -> str:
    if conc < 0.25:
        return "低"
    elif conc <= 0.4:
        return "中"
    return "高"


# ───────────── 生成辅助列 ─────────────

def add_helper_columns(df: pd.DataFrame, t: dict, s_threshold: int, a_threshold: int, stage: str,
                       params: dict | None = None) -> pd.DataFrame:
    col_cvr = get_col(df, ["24小时转化率", "转化率", "CVR"])
    t["col_cvr"] = col_cvr

    levels, reasons, match_types, competitions, pd_priorities, notes = [], [], [], [], [], []

    for _, row in df.iterrows():
        level, reason, match, pd_pri = classify_word(row, t, s_threshold, a_threshold, stage, params)
        try:
            conc = float(row.get(t["col_conc"], 0.5) or 0.5)
        except (ValueError, TypeError):
            conc = 0.5
        comp = get_competition(conc)

        levels.append(level)
        reasons.append(reason)
        match_types.append(match)
        competitions.append(comp)
        pd_priorities.append(pd_pri)
        notes.append("")  # 备注留空供人工复查

    df = df.copy()
    df["词级别"]       = levels
    df["分级理由"]      = reasons
    df["匹配方式建议"]  = match_types
    df["竞争密度"]      = competitions
    df["PD加码优先级"]  = pd_priorities
    df["备注"]          = notes
    return df


# ───────────── 生成分层汇总Sheet ─────────────

def build_summary_sheet(wb: openpyxl.Workbook, df: pd.DataFrame, t: dict,
                         s_threshold: int, a_threshold: int, stage: str):
    if "分层汇总" in wb.sheetnames:
        del wb["分层汇总"]
    ws = wb.create_sheet("分层汇总")
    ws.sheet_view.showGridLines = False

    head_fill  = PatternFill(start_color="1F4E79", end_color="1F4E79", fill_type="solid")
    head_font  = Font(color="FFFFFF", bold=True, name="Microsoft YaHei", size=11)
    title_font = Font(bold=True, name="Microsoft YaHei", size=13)
    body_font  = Font(name="Microsoft YaHei", size=10)
    alt_fill   = PatternFill(start_color="EBF3FB", end_color="EBF3FB", fill_type="solid")

    row = 1

    # ── 基准数据说明 ──
    ws.cell(row, 1, "📊 分层基准数据说明").font = title_font
    row += 1
    basis_info = [
        ("S级搜索量门槛", f"{s_threshold:,}（高相关词前3%）"),
        ("A级搜索量门槛", f"{a_threshold:,}（高相关词前15%）"),
        ("竞价中位数",    f"${t['bid_median']:.2f}"),
        ("集中度中位数",  f"{t['conc_median']:.3f}"),
        ("产品阶段",      stage),
    ]
    for k, v in basis_info:
        ws.cell(row, 1, k).font  = Font(bold=True, name="Microsoft YaHei", size=10)
        ws.cell(row, 2, v).font  = body_font
        row += 1
    row += 1

    # ── 各级别统计表 ──
    ws.cell(row, 1, "📈 各级别词数量统计").font = title_font
    row += 1
    stat_headers = ["词级别", "词数量", "占比", "平均周搜索量", "平均竞价中值", "平均集中度"]
    for ci, h in enumerate(stat_headers, 1):
        c = ws.cell(row, ci, h)
        c.fill, c.font, c.border = head_fill, head_font, THIN_BORDER
        c.alignment = Alignment(horizontal="center")
    row += 1

    col_search = t["col_search"]
    col_bid    = t["col_bid"]
    col_conc   = t["col_conc"]
    total = len(df)

    stat_start_row = row
    for level in ["S级", "A级", "B级", "C级"]:
        sub = df[df["词级别"] == level]
        cnt = len(sub)
        pct = f"{cnt/total*100:.1f}%" if total else "0%"
        avg_search = f"{sub[col_search].apply(pd.to_numeric, errors='coerce').mean():.0f}" if cnt else "0"
        avg_bid    = f"${sub[col_bid].apply(pd.to_numeric, errors='coerce').mean():.2f}" if cnt and col_bid else "N/A"
        avg_conc   = f"{sub[col_conc].apply(pd.to_numeric, errors='coerce').mean():.3f}" if cnt and col_conc else "N/A"
        fill_color = {"S级": "C6EFCE", "A级": "FFEB9C", "B级": "DDEBF7", "C级": "FCE4D6"}.get(level, "FFFFFF")
        row_fill = PatternFill(start_color=fill_color, end_color=fill_color, fill_type="solid")
        vals = [level, cnt, pct, avg_search, avg_bid, avg_conc]
        for ci, v in enumerate(vals, 1):
            c = ws.cell(row, ci, v)
            c.fill, c.border = row_fill, THIN_BORDER
            c.font = Font(name="Microsoft YaHei", size=10)
            c.alignment = Alignment(horizontal="center")
        row += 1
    stat_end_row = row - 1
    row += 1

    # ── S+A级重点词列表 ──
    ws.cell(row, 1, "🎯 S级+A级重点词完整列表（按搜索量排序）").font = title_font
    row += 1
    list_headers = ["词级别", "关键词", "中文翻译", "周搜索量", "竞价中值", "集中度", "匹配建议", "PD优先级", "分级理由"]
    for ci, h in enumerate(list_headers, 1):
        c = ws.cell(row, ci, h)
        c.fill, c.font, c.border = head_fill, head_font, THIN_BORDER
        c.alignment = Alignment(horizontal="center")
    row += 1

    col_kw   = get_col(df, ["关键词", "Keyword"])
    col_trans = get_col(df, ["中文翻译", "翻译", "Translation"])
    sa_df = df[df["词级别"].isin(["S级", "A级"])].copy()
    sa_df[col_search] = pd.to_numeric(sa_df[col_search], errors="coerce")
    sa_df = sa_df.sort_values(col_search, ascending=False)

    for i, (_, r) in enumerate(sa_df.iterrows()):
        fill = PatternFill(start_color="EBF3FB", end_color="EBF3FB", fill_type="solid") if i % 2 == 0 else None
        level = r.get("词级别", "")
        level_color = "C6EFCE" if level == "S级" else "FFEB9C"
        level_fill = PatternFill(start_color=level_color, end_color=level_color, fill_type="solid")
        vals = [
            r.get("词级别", ""),
            r.get(col_kw, "") if col_kw else "",
            r.get(col_trans, "") if col_trans else "",
            r.get(col_search, ""),
            r.get(col_bid, "") if col_bid else "",
            r.get(col_conc, "") if col_conc else "",
            r.get("匹配方式建议", ""),
            r.get("PD加码优先级", ""),
            r.get("分级理由", ""),
        ]
        for ci, v in enumerate(vals, 1):
            c = ws.cell(row, ci, v)
            c.border = THIN_BORDER
            c.font = Font(name="Microsoft YaHei", size=9)
            if ci == 1:
                c.fill = level_fill
            elif fill:
                c.fill = fill
        row += 1
    row += 1

    # ── PD前行动清单 ──
    ws.cell(row, 1, "📋 PD前行动清单").font = title_font
    row += 1
    actions = [
        ("S级词",  "单独开精准匹配广告活动，逐词建组，预算优先倾斜"),
        ("A级词",  "词组匹配，3-5个词一组"),
        ("B级词",  "广泛匹配统一测试组"),
        ("C级词",  "加入广告否定词列表"),
    ]
    for k, v in actions:
        ws.cell(row, 1, k).font = Font(bold=True, name="Microsoft YaHei", size=10)
        ws.cell(row, 2, v).font = body_font
        row += 1
    row += 1

    # ── 竞争密度判断 ──
    sa_conc = pd.to_numeric(sa_df[col_conc], errors="coerce") if col_conc else pd.Series(dtype=float)
    sa_bid  = pd.to_numeric(sa_df[col_bid],  errors="coerce") if col_bid  else pd.Series(dtype=float)
    avg_conc_val = sa_conc.mean() if len(sa_conc) > 0 else 0
    avg_bid_val  = sa_bid.mean()  if len(sa_bid)  > 0 else 0

    if avg_conc_val > 0.4 or avg_bid_val > 2.0:
        density_label = "竞争激烈"
    elif avg_conc_val < 0.25 and avg_bid_val < 1.0:
        density_label = "竞争温和"
    else:
        density_label = "竞争适中"

    ws.cell(row, 1, "🏁 类目竞争密度判断").font = title_font
    row += 1
    ws.cell(row, 1, f"结论：{density_label}").font = Font(bold=True, name="Microsoft YaHei", size=11,
        color="FF0000" if density_label == "竞争激烈" else ("00B050" if density_label == "竞争温和" else "000000"))
    row += 1
    ws.cell(row, 1, f"依据：S+A级平均集中度={avg_conc_val:.3f}，平均竞价中值=${avg_bid_val:.2f}").font = body_font
    row += 2

    # ── 嵌入openpyxl图表 ──
    counts = [len(df[df["词级别"] == lvl]) for lvl in ["S级", "A级", "B级", "C级"]]
    chart_data_row = row
    for i, (lvl, cnt) in enumerate(zip(["S级", "A级", "B级", "C级"], counts)):
        ws.cell(chart_data_row + i, 10, lvl)
        ws.cell(chart_data_row + i, 11, cnt)

    bar_chart = BarChart()
    bar_chart.type = "col"
    bar_chart.title = "各级别词数量分布"
    bar_chart.y_axis.title = "词数量"
    bar_chart.x_axis.title = "词级别"
    bar_chart.width = 16
    bar_chart.height = 10
    data_ref   = Reference(ws, min_col=11, min_row=chart_data_row, max_row=chart_data_row + 3)
    cats_ref   = Reference(ws, min_col=10, min_row=chart_data_row, max_row=chart_data_row + 3)
    bar_series = Series(data_ref, title="词数量")
    bar_chart.series.append(bar_series)
    bar_chart.set_categories(cats_ref)
    ws.add_chart(bar_chart, f"A{row}")

    # 图二：S+A级 搜索量 vs 集中度 散点图
    scatter_row = row
    scatter_n = min(len(sa_df), 50)
    for i in range(scatter_n):
        r = sa_df.iloc[i]
        sv = pd.to_numeric(r.get(col_search), errors="coerce")
        conc = pd.to_numeric(r.get(col_conc), errors="coerce") if col_conc else None
        if pd.isna(sv) or pd.isna(conc):
            continue
        ws.cell(scatter_row + i, 13, float(sv))
        ws.cell(scatter_row + i, 14, float(conc))
    if scatter_n > 0:
        scatter = ScatterChart()
        scatter.title = "S+A级 搜索量 vs 集中度（左下=高价值优先打）"
        scatter.x_axis.title = "周搜索量"
        scatter.y_axis.title = "Top3集中度"
        scatter.style = 13
        scatter.width = 16
        scatter.height = 10
        xvalues = Reference(ws, min_col=13, min_row=scatter_row, max_row=scatter_row + scatter_n - 1)
        yvalues = Reference(ws, min_col=14, min_row=scatter_row, max_row=scatter_row + scatter_n - 1)
        scatter.series.append(Series(yvalues, xvalues, title="S+A级词"))
        ws.add_chart(scatter, f"J{row}")

    row += 18

    # 调整列宽
    for ci, w in enumerate([12, 18, 18, 14, 12, 12, 14, 16, 45], 1):
        ws.column_dimensions[openpyxl.utils.get_column_letter(ci)].width = w

    return ws, density_label, avg_conc_val, avg_bid_val, stat_start_row, stat_end_row


# ───────────── 写入辅助列到原始Sheet ─────────────

def write_helper_cols_to_wb(wb: openpyxl.Workbook, df: pd.DataFrame,
                              src_path: str, t: dict):
    """把辅助列写回每个相关 Sheet"""
    helper_cols = ["词级别", "分级理由", "匹配方式建议", "竞争密度", "PD加码优先级", "备注"]
    sheet_meta = df.groupby("__sheet__")["__header_row__"].first().to_dict() if "__header_row__" in df.columns else {}

    for sheet_name in df["__sheet__"].unique():
        if sheet_name not in wb.sheetnames:
            continue
        ws = wb[sheet_name]
        max_col = ws.max_column
        header_row = int(sheet_meta.get(sheet_name, 2))

        for ci, hdr in enumerate(helper_cols, max_col + 1):
            c = ws.cell(header_row, ci, hdr)
            c.fill  = PatternFill(start_color="1F4E79", end_color="1F4E79", fill_type="solid")
            c.font  = Font(color="FFFFFF", bold=True, name="Microsoft YaHei")
            c.border = THIN_BORDER

        sheet_df = df[df["__sheet__"] == sheet_name].reset_index(drop=True)
        data_start = header_row + 1
        for ri, (_, row) in enumerate(sheet_df.iterrows()):
            excel_row = ri + data_start
            for ci, col in enumerate(helper_cols, max_col + 1):
                val = row.get(col, "")
                c = ws.cell(excel_row, ci, val)
                fill_color = {"S级": "C6EFCE", "A级": "FFEB9C", "B级": "DDEBF7", "C级": "FCE4D6"}.get(
                    row.get("词级别", ""), "FFFFFF")
                if col == "词级别":
                    c.fill = PatternFill(start_color=fill_color, end_color=fill_color, fill_type="solid")
                c.border = THIN_BORDER
                c.font = Font(name="Microsoft YaHei", size=9)


# ───────────── 生成Word报告 ─────────────

def build_t1_stats(df, t, s_threshold, a_threshold, stage, asin, date, params=None) -> dict:
    col_kw = get_col(df, ["关键词", "Keyword"])
    col_search = t["col_search"]
    level_counts = df["词级别"].value_counts().to_dict()
    sa_df = df[df["词级别"].isin(["S级", "A级"])].copy()
    sa_df[col_search] = pd.to_numeric(sa_df[col_search], errors="coerce")
    top_sa = sa_df.nlargest(10, col_search) if col_search else sa_df.head(10)
    top_list = []
    for _, r in top_sa.iterrows():
        top_list.append({
            "keyword": str(r.get(col_kw, "")) if col_kw else "",
            "level": str(r.get("词级别", "")),
            "search_volume": int(r.get(col_search, 0) or 0),
            "bid": float(r.get(t["col_bid"], 0) or 0) if t.get("col_bid") else None,
            "concentration": float(r.get(t["col_conc"], 0) or 0) if t.get("col_conc") else None,
            "match_type": str(r.get("匹配方式建议", "")),
            "pd_priority": str(r.get("PD加码优先级", "")),
        })
    return {
        "report": "t1_keyword_research",
        "asin": asin,
        "date": date,
        "stage": stage,
        "product_type": (params or {}).get("product_type", ""),
        "params": {k: (params or {}).get(k) for k in (
            "product_type", "s_conc_max", "a_conc_max", "s_bid_mult", "a_bid_mult",
            "b_require_cvr", "s_percentile", "a_percentile",
        ) if (params or {}).get(k) is not None},
        "total_keywords": len(df),
        "thresholds": {
            "s_search_volume": s_threshold,
            "a_search_volume": a_threshold,
            "bid_median": t["bid_median"],
            "conc_median": t["conc_median"],
        },
        "level_counts": {k: int(v) for k, v in level_counts.items()},
        "top_sa_keywords": top_list,
        "insights_prompt": "见同目录 insights_t1.md，或由 Agent 根据此 JSON 撰写",
    }


def build_word_report(df: pd.DataFrame, t: dict, s_threshold: int, a_threshold: int,
                       stage: str, asin: str, date: str, output_dir: str,
                       density_label: str, avg_conc: float, avg_bid: float,
                       insights: str = ""):
    if not HAS_DOCX:
        print("⚠️ 未安装 python-docx，跳过 Word。pip install python-docx", file=sys.stderr)
        return None

    doc = Document()
    setup_doc_styles(doc)
    add_doc_title(doc, f"{asin} 关键词调研分析报告", f"生成日期：{date}  |  产品阶段：{stage}")

    col_kw = get_col(df, ["关键词", "Keyword"])
    col_search = t["col_search"]
    col_bid = t["col_bid"]
    col_conc = t["col_conc"]

    sa_df = df[df["词级别"].isin(["S级", "A级"])].copy()
    sa_df[col_search] = pd.to_numeric(sa_df[col_search], errors="coerce")
    sa_df = sa_df.sort_values(col_search, ascending=False)

    level_counts = df["词级别"].value_counts()
    s_cnt = level_counts.get("S级", 0)
    a_cnt = level_counts.get("A级", 0)
    b_cnt = level_counts.get("B级", 0)
    c_cnt = level_counts.get("C级", 0)

    doc.add_heading("一、执行摘要", level=1)
    sa_top3 = sa_df.head(3)[col_kw].astype(str).tolist() if col_kw and len(sa_df) > 0 else []
    sa_top3_str = "、".join(sa_top3) if sa_top3 else "（无）"
    sa_avg_sv = pd.to_numeric(sa_df[col_search], errors="coerce").mean() if col_search and len(sa_df) > 0 else 0
    stage_advice = {
        "新品期": "新品期建议优先聚焦S级词用中等预算验证转化，A级词组匹配辅助。",
        "成长期": "成长期可加大S+A级词预算，同步布局B级词扩展流量。",
        "成熟期": "成熟期重点防守S级词搜索排名，C级词严格否词。",
    }.get(stage, "")
    add_para(doc,
        f"ASIN {asin}，{stage}阶段。共处理 {len(df)} 个关键词。"
        f"分层：S级 {s_cnt}、A级 {a_cnt}、B级 {b_cnt}、C级 {c_cnt}。"
        f"类目竞争：{density_label}（S+A平均集中度 {avg_conc:.3f}，平均竞价 ${avg_bid:.2f}）。"
        f"S级门槛≥{s_threshold:,}；A级≥{a_threshold:,}。"
        f"S+A TOP3：{sa_top3_str}，均值搜索量约 {sa_avg_sv:.0f}。{stage_advice}"
    )

    doc.add_heading("二、策略洞察与 PD 备战建议", level=1)
    render_insights_section(doc, insights, heading="")

    doc.add_heading("三、分层基准说明", level=1)
    basis_table = doc.add_table(rows=6, cols=2)
    basis_table.style = "Table Grid"
    for i, (k, v) in enumerate([
        ("S级搜索量门槛", f"{s_threshold:,}（高相关词周搜索量前3%）"),
        ("A级搜索量门槛", f"{a_threshold:,}（高相关词周搜索量前15%）"),
        ("竞价中位数", f"${t['bid_median']:.2f}"),
        ("集中度中位数", f"{t['conc_median']:.3f}"),
        ("S级竞价倍数", "×1.2"),
        ("S级集中度上限", "0.30"),
    ]):
        set_cell_text(basis_table.rows[i].cells[0], k)
        set_cell_text(basis_table.rows[i].cells[1], v)
    apply_font_to_table(basis_table)

    doc.add_heading("四、分层结果", level=1)
    result_table = doc.add_table(rows=5, cols=3)
    result_table.style = "Table Grid"
    for ci, h in enumerate(["词级别", "词数量", "代表词举例（3个）"]):
        set_cell_text(result_table.rows[0].cells[ci], h)
    for li, level in enumerate(["S级", "A级", "B级", "C级"], 1):
        sub = df[df["词级别"] == level]
        examples = "、".join(sub.head(3)[col_kw].astype(str).tolist()) if col_kw and len(sub) > 0 else ""
        set_cell_text(result_table.rows[li].cells[0], level)
        set_cell_text(result_table.rows[li].cells[1], str(len(sub)))
        set_cell_text(result_table.rows[li].cells[2], examples)
    apply_font_to_table(result_table)

    doc.add_heading("五、TOP20重点词解读", level=1)
    top20 = sa_df.head(20)
    if len(top20) > 0:
        t20 = doc.add_table(rows=len(top20) + 1, cols=6)
        t20.style = "Table Grid"
        for ci, h in enumerate(["关键词", "级别", "周搜索量", "竞价中值", "集中度", "PD优先级"]):
            set_cell_text(t20.rows[0].cells[ci], h)
        for ri, (_, r) in enumerate(top20.iterrows(), 1):
            set_cell_text(t20.rows[ri].cells[0], str(r.get(col_kw, "")) if col_kw else "")
            set_cell_text(t20.rows[ri].cells[1], str(r.get("词级别", "")))
            set_cell_text(t20.rows[ri].cells[2], str(int(r.get(col_search, 0) or 0)))
            set_cell_text(t20.rows[ri].cells[3], str(r.get(col_bid, "")) if col_bid else "")
            set_cell_text(t20.rows[ri].cells[4], str(r.get(col_conc, "")) if col_conc else "")
            set_cell_text(t20.rows[ri].cells[5], str(r.get("PD加码优先级", "")))
        apply_font_to_table(t20)

    if HAS_MPL:
        chart_paths = _generate_charts_table1(df, col_kw, col_search, col_conc, col_bid, output_dir, asin, date)
        items = []
        if len(chart_paths) >= 1:
            items.append(("图1 · 柱形图：S/A/B/C 各级别关键词数量分布", chart_paths[0]))
        if len(chart_paths) >= 2:
            items.append((
                "图2 · 气泡散点图：S+A 级词 周搜索量 vs Top3 集中度（气泡大小=竞价中值；"
                "绿=S级，黄=A级；左下=高价值优先打）",
                chart_paths[1],
            ))
        add_charts_section(doc, items, section_title="六、可视化图表", level=1)

    finalize_doc_fonts(doc)
    out_path = os.path.join(output_dir, f"{asin}_关键词调研分析报告_{date}.docx")
    try:
        doc.save(out_path)
    except PermissionError:
        out_path = os.path.join(output_dir, f"{asin}_关键词调研分析报告_{date}_new.docx")
        doc.save(out_path)
        print(f"WARN: 原 Word 被占用，已另存为 {out_path}", file=sys.stderr)
    print(f"OK Word: {out_path}")
    return out_path


def _generate_charts_table1(df, col_kw, col_search, col_conc, col_bid, output_dir, asin, date):
    paths = []
    # 图1：各级别词数量柱状图
    try:
        counts = [len(df[df["词级别"] == lvl]) for lvl in ["S级", "A级", "B级", "C级"]]
        fig, ax = plt.subplots(figsize=(8, 5))
        colors = ["#00B050", "#FFC000", "#4472C4", "#FF6B6B"]
        bars = ax.bar(["S级", "A级", "B级", "C级"], counts, color=colors, width=0.5)
        for bar, cnt in zip(bars, counts):
            ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.5,
                    str(cnt), ha="center", va="bottom", fontsize=12, fontweight="bold")
        ax.set_title("各级别词数量分布", fontsize=14, fontweight="bold")
        ax.set_xlabel("词级别")
        ax.set_ylabel("词数量")
        ax.spines[["top", "right"]].set_visible(False)
        path1 = chart_path(output_dir, f"chart1_t1_{asin}_{date}.png")
        fig.savefig(path1, dpi=150, bbox_inches="tight")
        plt.close(fig)
        paths.append(path1)
    except Exception as e:
        print(f"图1生成失败: {e}", file=sys.stderr)

    # 图2：S+A级词搜索量vs集中度气泡散点图
    try:
        sa = df[df["词级别"].isin(["S级", "A级"])].copy()
        sa[col_search] = pd.to_numeric(sa[col_search], errors="coerce")
        if col_conc:
            sa[col_conc] = pd.to_numeric(sa[col_conc], errors="coerce")
        if col_bid:
            sa[col_bid] = pd.to_numeric(sa[col_bid], errors="coerce").fillna(0.5)
        else:
            sa["__bid"] = 0.5
            col_bid = "__bid"
        sa = sa.dropna(subset=[col_search])
        if len(sa) > 0 and col_conc:
            fig, ax = plt.subplots(figsize=(9, 6))
            s_mask = sa["词级别"] == "S级"
            a_mask = sa["词级别"] == "A级"
            s_sizes = (sa.loc[s_mask, col_bid].clip(lower=0.3) * 40).values
            a_sizes = (sa.loc[a_mask, col_bid].clip(lower=0.3) * 35).values
            ax.scatter(
                sa.loc[s_mask, col_search], sa.loc[s_mask, col_conc],
                c="#00B050", label="S级", alpha=0.7, s=s_sizes, edgecolors="white", linewidths=0.4,
            )
            ax.scatter(
                sa.loc[a_mask, col_search], sa.loc[a_mask, col_conc],
                c="#FFC000", label="A级", alpha=0.7, s=a_sizes, marker="^",
                edgecolors="white", linewidths=0.4,
            )
            ax.axhline(y=0.30, color="red", linestyle="--", alpha=0.5, label="集中度0.30线")
            ax.set_xlabel("周搜索量（左下角=高价值优先打）")
            ax.set_ylabel("Top3集中度")
            ax.set_title("S+A级词 搜索量 vs 集中度（气泡大小=竞价中值）", fontsize=13, fontweight="bold")
            ax.legend()
            ax.spines[["top", "right"]].set_visible(False)
            path2 = chart_path(output_dir, f"chart2_t1_{asin}_{date}.png")
            fig.savefig(path2, dpi=150, bbox_inches="tight")
            plt.close(fig)
            paths.append(path2)
    except Exception as e:
        print(f"图2生成失败: {e}", file=sys.stderr)

    return paths


# ───────────── 主函数 ─────────────

def main():
    from datetime import date as _date, datetime
    from threshold_presets import (
        add_product_type_arg, resolve_thresholds, format_params_summary,
        load_config_product_type, collect_overrides_from_args,
    )

    parser = argparse.ArgumentParser(description="处理Sif关键词调研表（表1）")
    parser.add_argument("--mode",          choices=["compute-thresholds", "process"], default="process")
    parser.add_argument("--input",         required=True, help="原始Excel文件路径")
    parser.add_argument("--output-dir",    help="输出目录（process模式必填）")
    parser.add_argument("--asin",          help="ASIN（process模式必填）")
    parser.add_argument("--date",          default=_date.today().strftime("%Y%m%d"),
                        help="日期YYYYMMDD，默认今天")
    parser.add_argument("--stage",         default="新品期", choices=["新品期", "成长期", "成熟期"])
    add_product_type_arg(parser, default="非标品")
    parser.add_argument("--s-threshold",   type=int, default=0, help="覆盖S级搜索量门槛（0=自动计算）")
    parser.add_argument("--a-threshold",   type=int, default=0, help="覆盖A级搜索量门槛（0=自动计算）")
    parser.add_argument("--s-conc-max",    type=float, default=None, help="覆盖S级集中度上限")
    parser.add_argument("--a-conc-max",    type=float, default=None, help="覆盖A级集中度上限")
    parser.add_argument("--insights",      default="", help="AI 分析段（内联）")
    parser.add_argument("--insights-file", default="", help="AI 分析 markdown 文件路径")
    parser.add_argument("--skip-word",     action="store_true", help="跳过 Word，仅 Excel+stats+charts")
    args = parser.parse_args()
    args.insights = load_insights(args.insights, args.insights_file)

    overrides = collect_overrides_from_args(args)
    params = resolve_thresholds(args.product_type, args.stage, overrides)

    if not os.path.exists(args.input):
        print(f"❌ 找不到输入文件：{args.input}", file=sys.stderr)
        sys.exit(1)

    print(f"📂 读取文件：{args.input}")
    df = load_excel(args.input)
    print(f"📊 共读取 {len(df)} 行数据")

    t = compute_thresholds(df, params["s_percentile"], params["a_percentile"])

    s_threshold = params.get("s_threshold") or (args.s_threshold if args.s_threshold > 0 else t["s_threshold"])
    a_threshold = params.get("a_threshold") or (args.a_threshold if args.a_threshold > 0 else t["a_threshold"])

    print("\n" + "=" * 50)
    print("📐 分层基准数据（请确认后继续）：")
    print(f"  预设：{format_params_summary(params)}")
    print(f"  高相关词数量：{t['total_high']} / 总词数：{t['total_words']}")
    print(f"  S级搜索量门槛：{s_threshold:,}（高相关词 {params.get('s_percentile', 97)} 分位数，非预设固定值）")
    print(f"  A级搜索量门槛：{a_threshold:,}（高相关词 {params.get('a_percentile', 85)} 分位数）")
    print(f"  竞价中位数：${t['bid_median']:.2f}")
    print(f"  集中度中位数：{t['conc_median']:.3f}")
    print("=" * 50)

    if args.mode == "compute-thresholds":
        print(json.dumps({
            "product_type": params["product_type"],
            "params": params,
            "s_threshold": s_threshold,
            "a_threshold": a_threshold,
            "bid_median": t["bid_median"],
            "conc_median": t["conc_median"],
            "total_high": t["total_high"],
            "total_words": t["total_words"],
        }, ensure_ascii=False, indent=2))
        return

    # process 模式
    if not args.output_dir or not args.asin or not args.date:
        print("❌ process模式需要 --output-dir、--asin、--date 参数", file=sys.stderr)
        sys.exit(1)

    rid = make_run_id(args.date)
    os.makedirs(args.output_dir, exist_ok=True)

    print("\n🔄 计算辅助列...")
    df = add_helper_columns(df, t, s_threshold, a_threshold, args.stage, params)

    level_counts = df["词级别"].value_counts()
    print(f"  S级：{level_counts.get('S级', 0)} | A级：{level_counts.get('A级', 0)} | B级：{level_counts.get('B级', 0)} | C级：{level_counts.get('C级', 0)}")

    # 复制原始文件并追加辅助列
    import shutil
    out_excel = os.path.join(args.output_dir, f"{args.asin}_关键词分层分析_{rid}.xlsx")
    shutil.copy2(args.input, out_excel)

    print(f"\n📝 写入辅助列到Excel...")
    wb = openpyxl.load_workbook(out_excel)
    write_helper_cols_to_wb(wb, df, args.input, t)

    # 新建分层汇总Sheet
    print("📊 生成分层汇总Sheet...")
    ws_summary, density_label, avg_conc, avg_bid, _, _ = build_summary_sheet(
        wb, df, t, s_threshold, a_threshold, args.stage
    )

    # 保存Excel
    wb.save(out_excel)
    print(f"✅ Excel已生成：{out_excel}")

    stats = build_t1_stats(df, t, s_threshold, a_threshold, args.stage, args.asin, args.date, params)
    stats["run_id"] = rid
    sp = save_stats(stats, stats_path(args.output_dir, args.asin, "t1", rid))
    print(f"📊 Stats JSON：{sp}")

    if args.skip_word:
        print("\n" + "=" * 60)
        print("📝 [Agent] 读取 stats JSON，写 insights_t1.md（分析段，不含图表）")
        print(f"  stats：{sp}")
        print(f"  insights 写入：{insights_path(args.output_dir, 't1')}")
        print("  然后带 --insights-file 重建 Word")
        print("=" * 60)
    else:
        print("\n📄 生成 Word 分析报告...")
        build_word_report(df, t, s_threshold, a_threshold, args.stage,
                          args.asin, rid, args.output_dir,
                          density_label, avg_conc, avg_bid,
                          insights=args.insights)
        cleanup_intermediate_files(args.output_dir, args.asin, rid, "t1")

    print(f"\n✅ 表1处理完成！输出目录：{args.output_dir}")


if __name__ == "__main__":
    main()
